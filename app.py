import sounddevice as sd
from openai import OpenAI
import numpy as np
import soundfile as sf
import os
import subprocess
from dotenv import load_dotenv
import time
from queue import Queue
import asyncio
import concurrent.futures
import logging
from flask import Flask, request, jsonify, render_template, redirect, send_from_directory, Response, session, url_for, flash
from flask_cors import CORS
from datetime import datetime, timedelta
import json
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from asgiref.wsgi import WsgiToAsgi
import requests
import speech_recognition as sr
from concurrent.futures import ThreadPoolExecutor
import threading
import random
import string
import hashlib
import sqlite3
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import secrets
import aiohttp
import queue
import playsound
import wave
from functools import wraps
import re
import base64
from flask_session import Session
import pathlib
import traceback
import PyPDF2
from docx import Document
from werkzeug.utils import secure_filename
import openai

# .env dosyasını yükle
load_dotenv()

# Klasör kontrolleri
QUESTIONS_DIR = 'interview_questions'
os.makedirs(QUESTIONS_DIR, exist_ok=True)

# Logging seviyesini ayarla - DEBUG yerine INFO kullan
logging.basicConfig(level=logging.INFO)

# Flask ve async ayarları
app = Flask(__name__)
app.config['SESSION_TYPE'] = 'filesystem'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=24)
app.config['SESSION_COOKIE_SECURE'] = False  # Development için False
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.secret_key = os.getenv('SECRET_KEY', 'duftech_secret_key_2024')
Session(app)

# CORS ayarları
CORS(app, supports_credentials=True)

# Logging ayarları
logger = logging.getLogger(__name__)

# OpenAI istemcisini başlat
openai_client = OpenAI()

# Açık erişime izin verilen route'lar
OPEN_ROUTES = ['login', 'join', 'static', 'interview', 'verify_code', 'submit_quiz', 'quiz_entry', 'check_quiz_code']

# Gerekli dizinleri oluştur
required_dirs = ['reports', 'temp', 'interview_questions', 'interviews']
for dir_name in required_dirs:
    os.makedirs(dir_name, exist_ok=True)

# E-posta ayarları
SMTP_SERVER = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
SMTP_PORT = int(os.getenv('SMTP_PORT', 587))
SMTP_USERNAME = os.getenv('SMTP_USERNAME')
SMTP_PASSWORD = os.getenv('SMTP_PASSWORD')
REPORT_SENDER = os.getenv('REPORT_SENDER')
REPORT_RECIPIENT = os.getenv('REPORT_RECIPIENT')

# Webhook URL'lerini güncelle
WEBHOOK_ADAY_URL = os.getenv('WEBHOOK_ADAY_URL')
WEBHOOK_RAPOR_URL = os.getenv('WEBHOOK_RAPOR_URL')

# Global değişkenler
interview_data = {}  # Mülakat verilerini saklamak için
SILENCE_THRESHOLD = 0.05  # Sessizlik eşik değeri
VOICE_THRESHOLD = 0.15   # Ses algılama eşik değeri
SILENCE_DURATION = 1500  # Sessizlik süresi (ms)
MIN_CONFIDENCE = 0.6     # Minimum güven skoru

# Paket tipleri
PACKAGE_TYPES = {
    'basic': {
        'name': 'Temel Paket',
        'features': ['interview', 'quiz'],
        'max_interviews': 5,
        'max_quizzes': 5,
        'price': 0
    },
    'professional': {
        'name': 'Profesyonel Paket',
        'features': ['interview', 'quiz', 'cv_analysis', 'reports'],
        'max_interviews': 20,
        'max_quizzes': 20,
        'price': 100
    },
    'enterprise': {
        'name': 'Kurumsal Paket',
        'features': ['interview', 'quiz', 'cv_analysis', 'reports', 'admin_panel'],
        'max_interviews': 100,
        'max_quizzes': 100,
        'price': 500
    }
}

class VoiceAssistant:
    def __init__(self):
        self.sample_rate = 16000
        self.channels = 1
        self.silence_threshold = 0.05
        self.silence_duration = 0.5
        self.audio_queue = Queue()
        self.is_recording = False
        self.executor = concurrent.futures.ThreadPoolExecutor(max_workers=4)

        # API ayarları
        openai_api_key = os.getenv('OPENAI_API_KEY')
        self.openai_client = OpenAI(api_key=openai_api_key)
        logger.info("OpenAI client başarıyla başlatıldı")

    def record_audio(self):
        """Ses kaydı yap"""
        logger.debug("Ses kaydı başlıyor...")
        audio_chunks = []
        silence_start = None
        self.is_recording = True
        
        # Ses algılama parametreleri
        SILENCE_THRESHOLD = 0.02  # Sessizlik eşiği
        MIN_AUDIO_LENGTH = 0.3    # Minimum ses süresi (saniye)
        MAX_SILENCE_LENGTH = 1.5  # Maksimum sessizlik süresi (saniye)
        BUFFER_SIZE = 1024        # Tampon boyutu
        
        try:
            with sd.InputStream(callback=self.audio_callback,
                              channels=self.channels,
                              samplerate=self.sample_rate,
                              blocksize=BUFFER_SIZE,
                              device=None):  # Varsayılan ses cihazı
                
                last_active_time = time.time()
                
                while self.is_recording:
                    try:
                        audio_chunk = self.audio_queue.get(timeout=0.1)
                        current_volume = np.max(np.abs(audio_chunk))
                        
                        # Ses algılandı
                        if current_volume > SILENCE_THRESHOLD:
                            last_active_time = time.time()
                            if silence_start is not None:
                                silence_duration = time.time() - silence_start
                                if silence_duration > MAX_SILENCE_LENGTH:
                                    audio_chunks = []  # Tampon temizlenir
                                silence_start = None
                            audio_chunks.append(audio_chunk)
                        # Sessizlik algılandı
                        else:
                            if silence_start is None:
                                silence_start = time.time()
                            elif time.time() - silence_start > MAX_SILENCE_LENGTH:
                                # Yeterli uzunlukta ses kaydı varsa bitir
                                if len(audio_chunks) * (BUFFER_SIZE / self.sample_rate) >= MIN_AUDIO_LENGTH:
                                    break
                            
                        # Uzun süre ses algılanmazsa kaydı durdur
                        if time.time() - last_active_time > 5:  # 5 saniye sessizlik
                            if len(audio_chunks) > 0:
                                break
                            else:
                                audio_chunks = []  # Tampon temizle
                                last_active_time = time.time()  # Zamanı sıfırla
                                
                    except queue.Empty:
                        continue
                        
        except Exception as e:
            logger.error(f"Ses kaydı hatası: {str(e)}")
            return None
            
        logger.debug("Ses kaydı tamamlandı")
        return np.concatenate(audio_chunks) if audio_chunks else None

    def audio_callback(self, indata, frames, time, status):
        """Ses verisi callback fonksiyonu"""
        if status:
            logger.warning(f"Ses kaydı durum: {status}")
        try:
            self.audio_queue.put(indata.copy())
        except queue.Full:
            self.audio_queue.get_nowait()  # En eski veriyi at
            self.audio_queue.put(indata.copy())

    def save_audio(self, recording, filename='temp_recording.wav'):
        logger.debug(f"Ses kaydı kaydediliyor: {filename}")
        sf.write(filename, recording, self.sample_rate)
        return filename

    async def transcribe_audio(self, audio_file):
        try:
            logger.debug("Ses tanıma başlıyor...")
            
            # WebM'den WAV'a dönüştür
            converted_file = 'temp/temp_converted.wav'
            ffmpeg_command = [
                'ffmpeg', '-y',
                '-i', audio_file,
                '-acodec', 'pcm_s16le',
                '-ac', '1',
                '-ar', '48000',
                converted_file
            ]
            
            try:
                subprocess.run(ffmpeg_command, check=True, capture_output=True)
                logger.info("WebM dosyası WAV formatına dönüştürüldü")
            except subprocess.CalledProcessError as e:
                logger.error(f"FFmpeg dönüştürme hatası: {e.stderr.decode()}")
                return None

            # OpenAI Whisper API kullanarak ses tanıma
            with open(converted_file, 'rb') as audio:
                transcript = self.openai_client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio,
                    language="tr"
                )

            # Geçici dosyaları temizle
            for temp_file in [converted_file]:
                if os.path.exists(temp_file):
                    try:
                        os.remove(temp_file)
                        logger.info(f"Geçici dosya silindi: {temp_file}")
                    except Exception as e:
                        logger.warning(f"Geçici dosya silme hatası: {str(e)}")
            
            return transcript.text.strip()
            
        except Exception as e:
            logger.error(f"Ses tanıma hatası: {str(e)}")
            return None

    async def generate_and_play_speech(self, text):
        """OpenAI TTS ile metni sese çevirir ve oynatır"""
        try:
            # Tuple ise ilk elemanı al, değilse direkt metni kullan
            if isinstance(text, tuple):
                text = text[0]
                
            logger.info(f"OpenAI TTS ile ses oluşturuluyor... Metin: {text[:50]}...")
            
            # Temp klasörünü kontrol et
            if not os.path.exists('temp'):
                os.makedirs('temp')
                
            speech_file_path = "temp/temp_speech.mp3"
            
            # OpenAI TTS API çağrısı
            response = self.openai_client.audio.speech.create(
                model="tts-1",
                voice="nova",
                input=text
            )
            
            # Ses dosyasını kaydet
            response.stream_to_file(speech_file_path)
            
            # Sesi oynat
            playsound.playsound(speech_file_path)
            
            # Geçici dosyayı sil
            if os.path.exists(speech_file_path):
                os.remove(speech_file_path)
                logger.info("Geçici ses dosyası silindi")
                
            return True
                
        except Exception as e:
            logger.error(f"OpenAI TTS API hatası: {e}")
            return False

   

class InterviewAssistant(VoiceAssistant):
    def __init__(self):
        super().__init__()
        self.code = None
        self.candidate_name = ""
        self.position = ""
        self.requirements = []
        self.custom_questions = []
        self.cv_summary = ""
        self.pre_info = ""
        self.conversation_history = []
        self.sentiment_scores = []
        self.metrics = {
            "iletisim_puani": 0,
            "ozguven_puani": 0,
            "teknik_bilgi": 0,
            "genel_puan": 0
        }
        self.start_time = datetime.now()
        self.current_question_index = 0
        self.interview_questions = []

    def set_interview_details(self, code):
        """Mülakat detaylarını JSON dosyasından ayarla"""
        try:
            # JSON dosyasından mülakat bilgilerini oku
            json_path = os.path.join('interviews', f'{code}.json')
            if not os.path.exists(json_path):
                raise ValueError(f"Mülakat dosyası bulunamadı: {code}")

            with open(json_path, 'r', encoding='utf-8') as f:
                interview_data = json.load(f)

            # Mülakat bilgilerini ayarla
            self.code = code
            self.candidate_name = interview_data.get("candidate_name", "")
            self.position = interview_data.get("position", "")
            self.requirements = interview_data.get("requirements", [])
            self.custom_questions = interview_data.get("questions", [])
            self.cv_summary = interview_data.get("cv_summary", "")
            self.pre_info = interview_data.get("pre_info", "")

            # CV verilerini kaydet (eğer varsa)
            self.cv_data = interview_data.get("cv_data", {})
            
            # CV'den beceri ve deneyimleri çıkar (sorular için kullanılacak)
            if self.cv_data:
                self.skills = self.cv_data.get("skills", [])
                self.experience = self.cv_data.get("experience", [])
                self.education = self.cv_data.get("education", [])
                self.profile_summary = self.cv_data.get("profile_summary", "")
                
                # Log - CV başarıyla yüklendi
                logger.info(f"CV verileri başarıyla yüklendi - Mülakat Kodu: {code}")
            else:
                logger.info(f"CV verisi bulunamadı - Mülakat Kodu: {code}")

            # Soruları hazırla
            self._prepare_interview_questions()

            logger.info(f"Mülakat detayları ayarlandı: {code}")
            return True

        except Exception as e:
            logger.error(f"Mülakat detayları ayarlama hatası: {str(e)}")
            raise

    def _prepare_interview_questions(self):
        """Özelleştirilmiş mülakat sorularını hazırla"""
        try:
            # Özel sorular varsa, bunları kullan
            if self.custom_questions and len(self.custom_questions) > 0:
                self.interview_questions = self.custom_questions
                return
            
            # Özel sorular yoksa ve CV verileri varsa, CV'ye özel sorular oluştur
            if hasattr(self, 'cv_data') and self.cv_data:
                cv_based_questions = []
                
                # Becerilerden soru oluştur
                if hasattr(self, 'skills') and self.skills:
                    # En önemli 2-3 beceriyi seç
                    top_skills = self.skills[:3]
                    for skill in top_skills:
                        cv_based_questions.append(f"{skill} konusundaki deneyimlerinizi detaylandırabilir misiniz?")
                
                # Deneyimlerden soru oluştur
                if hasattr(self, 'experience') and self.experience and len(self.experience) > 0:
                    # En son deneyimden soru oluştur
                    latest_exp = self.experience[0]
                    company = latest_exp.get('company', '')
                    title = latest_exp.get('title', '')
                    
                    if company and title:
                        cv_based_questions.append(f"{company}'da {title} olarak çalışırken karşılaştığınız en büyük zorluk neydi ve bunu nasıl aştınız?")
                    
                    # Birden fazla deneyim varsa, karşılaştırma sorusu sor
                    if len(self.experience) > 1:
                        cv_based_questions.append("Farklı şirketlerdeki deneyimlerinizi karşılaştırdığınızda, size en çok katkı sağlayan hangisiydi ve neden?")
                
                # Eğitimden soru oluştur
                if hasattr(self, 'education') and self.education and len(self.education) > 0:
                    education = self.education[0]
                    degree = education.get('degree', '')
                    institution = education.get('institution', '')
                    
                    if degree and institution:
                        cv_based_questions.append(f"{institution}'da aldığınız {degree} eğitiminin şu anki kariyerinize nasıl katkı sağladığını düşünüyorsunuz?")
                
                # Genel sorular ekle
                general_questions = [
                    "Bize biraz kendinizden ve kariyer hedeflerinizden bahseder misiniz?",
                    f"{self.position} pozisyonuna neden başvurdunuz?",
                    "Takım çalışması hakkında neler düşünüyorsunuz?",
                    "Baskı altında çalışma konusunda nasıl bir yaklaşımınız var?",
                    "Bu pozisyonda 1 yıl sonra kendinizi nerede görüyorsunuz?"
                ]
                
                # CV'den oluşturulan sorular ve genel soruları birleştir
                self.interview_questions = cv_based_questions + general_questions
                
                # En fazla 10 soru olacak şekilde sınırla
                if len(self.interview_questions) > 10:
                    self.interview_questions = self.interview_questions[:10]
                
                logger.info(f"CV verilerine dayalı {len(cv_based_questions)} adet soru oluşturuldu")
            else:
                # Varsayılan sorular
                self.interview_questions = [
                    "Bize biraz kendinizden bahseder misiniz?",
                    "Bu pozisyona neden başvurdunuz?",
                    "Daha önce benzer rollerde çalıştınız mı?",
                    "Teknik becerileriniz hakkında bilgi verebilir misiniz?",
                    "Takım çalışması hakkında neler düşünüyorsunuz?",
                    "Zorlu bir iş durumu ile karşılaştığınızda nasıl bir yaklaşım sergilersiniz?",
                    "Baskı altında çalışma konusunda nasıl bir yaklaşımınız var?",
                    "Güçlü ve geliştirilmesi gereken yönleriniz nelerdir?",
                    "Bu pozisyonda 1 yıl sonra kendinizi nerede görüyorsunuz?",
                    "Bizim için size sormak istediğiniz herhangi bir soru var mı?"
                ]
            
        except Exception as e:
            logger.error(f"Mülakat soruları hazırlama hatası: {str(e)}")
            # Hata durumunda varsayılan sorular
            self.interview_questions = [
                "Bize biraz kendinizden bahseder misiniz?",
                "Bu pozisyona neden başvurdunuz?",
                "Daha önce benzer rollerde çalıştınız mı?",
                "Teknik becerileriniz hakkında bilgi verebilir misiniz?",
                "Ekip çalışması hakkında ne düşünüyorsunuz?"
            ]

    async def get_gpt_response(self, text):
        try:
            # Başlangıç kontrolü
            if not hasattr(self, 'interview_started'):
                if any(word in text.lower() for word in ["başlayalım", "başla", "hazırım", "başlayabiliriz"]):
                    self.interview_started = True
                    self.current_question_index = 0
                    
                    # CV bilgilerini kontrol et
                    cv_greeting = ""
                    if hasattr(self, 'cv_data') and self.cv_data:
                        # CV verilerinden kişiselleştirilmiş giriş metni oluştur
                        experience = self.cv_data.get('experience', [])
                        skills = self.cv_data.get('skills', [])
                        education = self.cv_data.get('education', [])
                        
                        if experience or skills:
                            cv_greeting = "\n\nÖzgeçmişinizi inceledim. "
                            
                            if experience and len(experience) > 0:
                                last_company = experience[0].get('company', '')
                                last_position = experience[0].get('title', '')
                                if last_company and last_position:
                                    cv_greeting += f"{last_company}'da {last_position} olarak çalışma deneyiminiz dikkatimi çekti. "
                            
                            if skills and len(skills) > 0:
                                top_skills = skills[:3]
                                cv_greeting += f"{', '.join(top_skills)} gibi becerilerinizin bu pozisyon için değerli olduğunu düşünüyorum."
                    
                    welcome_message = f"""Merhaba {self.candidate_name}, {self.position} pozisyonu için mülakatımıza hoş geldiniz.{cv_greeting}

Pozisyonun gerektirdiği yetkinlikler hakkında konuşacağız.

İlk olarak, {self.interview_questions[self.current_question_index]}"""
                    return welcome_message, False
                else:
                    return "Mülakata başlamak için 'Başlayalım' diyebilirsiniz.", False

            # Tekrar isteği kontrolü
            if any(phrase in text.lower() for phrase in ["tekrar", "anlamadım", "tekrarlar mısın", "bir daha söyler misin"]):
                return f"Tabii ki, soruyu tekrar ediyorum: {self.interview_questions[self.current_question_index]}", False

            # Kısa cevap kontrolü
            if len(text.split()) < 5:
                follow_up_responses = [
                    "Bu konu hakkında biraz daha detay verebilir misiniz?",
                    "İlginç, peki bu konuda başka neler söylemek istersiniz?",
                    "Deneyimlerinizden örnekler verebilir misiniz?",
                    "Bu konuyu biraz daha açar mısınız?"
                ]
                return random.choice(follow_up_responses), False

            # Bir sonraki soruya geçiş
            if self.current_question_index < len(self.interview_questions) - 1:
                self.current_question_index += 1
                
                # CV verilerine göre kişiselleştirilmiş geçiş cümleleri
                cv_reference = ""
                if hasattr(self, 'cv_data') and self.cv_data and random.random() < 0.3:  # %30 olasılıkla CV'ye atıf yap
                    experience = self.cv_data.get('experience', [])
                    skills = self.cv_data.get('skills', [])
                    
                    # Sonraki sorunun içeriğine göre CV'den ilgili bilgileri referans al
                    next_question = self.interview_questions[self.current_question_index].lower()
                    
                    if any(skill.lower() in next_question for skill in skills):
                        matching_skills = [skill for skill in skills if skill.lower() in next_question]
                        if matching_skills:
                            cv_reference = f" Özgeçmişinizde {matching_skills[0]} konusunda deneyiminiz olduğunu görmüştüm."
                    
                    elif any(keyword in next_question for keyword in ["deneyim", "tecrübe", "çalışma"]):
                        if experience and len(experience) > 0:
                            cv_reference = f" Bu soru, {experience[0].get('company', '')} şirketindeki deneyiminizle ilgili olabilir."
                
                transition_phrases = [
                    f"Anlıyorum, teşekkür ederim. Şimdi başka bir konuya geçelim.{cv_reference} {self.interview_questions[self.current_question_index]}",
                    f"Bu konudaki görüşleriniz için teşekkürler. Peki,{cv_reference} {self.interview_questions[self.current_question_index]}",
                    f"Güzel bir açıklama oldu. İsterseniz şimdi{cv_reference} {self.interview_questions[self.current_question_index]}",
                    f"Teşekkür ederim. Bir sonraki konumuz:{cv_reference} {self.interview_questions[self.current_question_index]}"
                ]
                return random.choice(transition_phrases), False
            else:
                closing_message = """Mülakat sorularımız tamamlandı. Paylaştığınız değerli bilgiler ve ayırdığınız zaman için teşekkür ederiz. 
                
Değerlendirme sonucunu en kısa sürede size ileteceğiz. İyi günler dilerim."""
                return closing_message, True

        except Exception as e:
            logger.error(f"GPT yanıtı alma hatası: {str(e)}")
            return "Üzgünüm, bir hata oluştu. Lütfen tekrar dener misiniz?", False

    def prepare_conversation_context(self, text):
        return [
            {"role": "system", "content": f"Sen bir mülakat uzmanısın. {self.position} pozisyonu için mülakat yapıyorsun."},
            {"role": "user", "content": text}
        ]

    async def _analyze_sentiment(self, text):
        """Adayın cevaplarını analiz et"""
        try:
            response = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: openai_client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "Verilen metni analiz et ve duygusal durumu değerlendir."},
                        {"role": "user", "content": text}
                    ],
                    temperature=0.7,
                    max_tokens=100
                )
            )
            
            sentiment = response.choices[0].message.content
            logger.info(f"Duygu analizi sonucu: {sentiment}")
            return sentiment
            
        except Exception as e:
            logger.error(f"Duygu analizi hatası: {str(e)}")
            return None

    async def process_interview_response(self, text):
        """Mülakat yanıtını işle ve değerlendir"""
        try:
            if not text:
                return None
                
            # Bir sonraki soruyu hazırla
            if self.current_question_index < len(self.interview_questions):
                next_question = self.interview_questions[self.current_question_index]
                self.current_question_index += 1
            else:
                next_question = "Mülakat sona erdi."
            
            # Konuşma geçmişini güncelle
            self.conversation_history.extend([
                {"role": "user", "content": text},
                {"role": "assistant", "content": next_question}
            ])
            
            return next_question
            
        except Exception as e:
            logger.error(f"Yanıt işleme hatası: {str(e)}")
            return None

    def generate_pdf_report(self):
        try:
            # PDF dosya yolu
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            pdf_path = os.path.join('reports', f"mulakat_raporu_{timestamp}.pdf")
            
            # PDF oluştur
            doc = SimpleDocTemplate(pdf_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Başlık
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30
            )
            story.append(Paragraph("Mülakat Raporu", title_style))
            story.append(Spacer(1, 12))
            
            # Mülakat bilgileri
            info_style = ParagraphStyle(
                'Info',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=6
            )
            story.append(Paragraph(f"Aday: {self.candidate_name if current_interview else 'Bilinmiyor'}", info_style))
            story.append(Paragraph(f"Pozisyon: {self.position if current_interview else 'Bilinmiyor'}", info_style))
            story.append(Paragraph(f"Tarih: {self.start_time.strftime('%d.%m.%Y %H:%M')}", info_style))
            story.append(Spacer(1, 20))
            
            # Performans metrikleri
            story.append(Paragraph("Performans Değerlendirmesi", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            metrics_data = [
                ["Metrik", "Puan"],
                ["İletişim Becerisi", f"{self.metrics['iletisim_puani']:.1f}"],
                ["Özgüven", f"{self.metrics['ozguven_puani']:.1f}"],
                ["Teknik Bilgi", f"{self.metrics['teknik_bilgi']:.1f}"],
                ["Genel Puan", f"{self.metrics['genel_puan']:.1f}"]
            ]
            
            t = Table(metrics_data, colWidths=[300, 100])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(t)
            story.append(Spacer(1, 20))
            
            # Konuşma geçmişi
            story.append(Paragraph("Mülakat Detayları", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            for entry in self.conversation_history:
                if entry["role"] == "user":
                    story.append(Paragraph(f"<b>Aday:</b> {entry['content']}", styles['Normal']))
                else:
                    story.append(Paragraph(f"<b>Mülakat Uzmanı:</b> {entry['content']}", styles['Normal']))
                story.append(Spacer(1, 12))
            
            # PDF oluştur
            doc.build(story)
            logger.info(f"PDF raporu başarıyla oluşturuldu: {pdf_path}")
            return pdf_path
            
        except Exception as e:
            logger.error(f"PDF rapor oluşturma hatası: {str(e)}")
            return None

    def send_report_email(self, pdf_path):
        try:
            # E-posta oluştur
            msg = MIMEMultipart()
            msg['From'] = REPORT_SENDER
            msg['To'] = REPORT_RECIPIENT
            msg['Subject'] = f"Mülakat Raporu - {self.candidate_name} - {self.position}"
            
            # E-posta metni
            body = f"""
            Merhaba,
            
            {self.candidate_name} adayı ile {self.position} pozisyonu için yapılan mülakat raporu ekte yer almaktadır.
            
            Mülakat Bilgileri:
            - Aday: {self.candidate_name}
            - Pozisyon: {self.position}
            - Tarih: {self.start_time.strftime('%d.%m.%Y %H:%M')}
            - Genel Puan: {self.metrics['genel_puan']:.1f}/100
            
            Detaylı değerlendirme için ekteki PDF dosyasını inceleyebilirsiniz.
            
            İyi çalışmalar,
            DUF Tech Mülakat Asistanı
            """
            msg.attach(MIMEText(body, 'plain'))
            
            # PDF ekle
            with open(pdf_path, "rb") as f:
                pdf = MIMEApplication(f.read(), _subtype="pdf")
                pdf.add_header('Content-Disposition', 'attachment', filename=os.path.basename(pdf_path))
                msg.attach(pdf)
            
            # E-postayı gönder
            with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
                server.starttls()
                server.login(SMTP_USERNAME, SMTP_PASSWORD)
                server.send_message(msg)

            # Webhook'a rapor gönder
            self.send_report_webhook(pdf_path)
            
            return True
        except Exception as e:
            print(f"E-posta gönderme hatası: {str(e)}")
            return False

    def send_report_webhook(self, pdf_path, evaluation_data):
        """Raporu webhook'a gönder"""
        try:
            # Mülakat verilerini hazırla
            interview_summary = {
                "mulakat_kodu": self.code,
                "aday_bilgileri": {
                    "isim": self.candidate_name,
                    "pozisyon": self.position,
                    "cv_ozeti": self.cv_summary,
                    "on_bilgi": self.pre_info,
                    "gereksinimler": self.requirements,
                    "tarih": self.start_time.isoformat()
                },
                "mulakat_metrikleri": {
                    "iletisim_puani": self.metrics["iletisim_puani"],
                    "ozguven_puani": self.metrics["ozguven_puani"],
                    "teknik_bilgi": self.metrics["teknik_bilgi"],
                    "genel_puan": self.metrics["genel_puan"]
                },
                "degerlendirme": evaluation_data,
                "konusma_akisi": self.conversation_history,
                "rapor_dosyasi": pdf_path,
                "olusturulma_tarihi": datetime.now().isoformat()
            }
            
            # JSON dosyasından webhook URL'sini kontrol et
            json_path = os.path.join('interviews', f'{self.code}.json')
            webhook_url = WEBHOOK_RAPOR_URL  # Varsayılan URL
            
            if os.path.exists(json_path):
                with open(json_path, 'r', encoding='utf-8') as f:
                    interview_json = json.load(f)
                    if 'webhook_rapor_url' in interview_json:
                        webhook_url = interview_json['webhook_rapor_url']
                        logger.info(f"JSON'dan webhook URL'si kullanılıyor: {webhook_url}")
                    else:
                        logger.info(f"Varsayılan webhook URL'si kullanılıyor: {webhook_url}")
            
            # WebhookRapor'a gönder
            response = requests.post(
                webhook_url,
                json=interview_summary,
                headers={
                    'Content-Type': 'application/json',
                    'X-Interview-Code': self.code
                }
            )
            
            if response.status_code == 200:
                logger.info(f"Rapor webhook'a başarıyla gönderildi: {self.code}")
                return True
            else:
                logger.error(f"Webhook hatası: {response.status_code} - {response.text}")
                return False
                
        except Exception as e:
            logger.error(f"Webhook gönderme hatası: {str(e)}")
            return False

def generate_interview_code():
    """Benzersiz bir mülakat kodu oluştur"""
    code = ''.join(secrets.choice(string.ascii_uppercase + string.digits) for _ in range(8))
    return code

# Login için gerekli decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('logged_in'):
            logger.debug("Oturum bulunamadı, login sayfasına yönlendiriliyor")
            return redirect(url_for('login'))
        logger.debug(f"Oturum doğrulandı: {session.get('user_id')}")
        return f(*args, **kwargs)
    return decorated_function

# Login bilgilerini .env'den al
ADMIN_EMAIL = os.getenv('ADMIN_EMAIL')
ADMIN_PASSWORD = os.getenv('ADMIN_PASSWORD')

# Login sayfası route'u
@app.route('/login', methods=['GET', 'POST'])
def login():
    """Kullanıcı giriş sayfası"""
    # Zaten giriş yapmış kullanıcıyı kontrol et
    if session.get('logged_in'):
        if session.get('is_admin'):
            return redirect(url_for('admin_dashboard'))
        else:
            return redirect(url_for('create_interview'))
    
    # POST işlemi - form gönderme
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        try:
            # Admin kullanıcı bilgilerini kontrol et
            admin_email = os.getenv('ADMIN_EMAIL')
            admin_password = os.getenv('ADMIN_PASSWORD')
            
            logger.debug(f"Beklenen admin: {admin_email}")
            
            # Admin girişi
            if username == admin_email and password == admin_password:
                session.permanent = True  # Oturumu kalıcı yap
                session['logged_in'] = True
                session['user_email'] = username
                session['is_admin'] = True
                session['user_name'] = 'Admin'
                logger.info(f"Admin girişi başarılı: {username}")
                return redirect(url_for('admin_dashboard'))
            
            # Normal kullanıcıları kontrol et (gerçek uygulamada veritabanından)
            # Örnek kullanıcı: kullanici@duftech.com / sifre123
            if username == 'kullanici@duftech.com' and password == 'sifre123':
                session.permanent = True  # Oturumu kalıcı yap
                session['logged_in'] = True
                session['user_email'] = username
                session['is_admin'] = False
                session['user_name'] = 'Kullanıcı Demo'
                logger.info(f"Kullanıcı girişi başarılı: {username}")
                return redirect(url_for('create_interview'))
            
            # Giriş başarısız
            logger.warning(f"Başarısız giriş denemesi: {username}")
            return render_template('login.html', error="Geçersiz kullanıcı adı veya şifre")
            
        except Exception as e:
            logger.error(f"Giriş işlemi hatası: {str(e)}")
            return render_template('login.html', error="Giriş işlemi sırasında bir hata oluştu")
    
    # GET işlemi - form görüntüleme
    return render_template('login.html')

@app.route('/logout')
def logout():
    """Çıkış yapma"""
    session.pop('logged_in', None)
    session.pop('user_email', None)
    session.pop('is_admin', None)
    session.pop('user_name', None)
    return redirect(url_for('login'))

@app.route('/')
def home():
    """Ana sayfa"""
    # Kullanıcı giriş kontrolü
    if session.get('logged_in'):
        # Admin kullanıcıyı admin paneline yönlendir
        if session.get('is_admin'):
            return redirect(url_for('admin_dashboard'))
        # Normal kullanıcıyı mülakat oluşturma sayfasına yönlendir
        return redirect(url_for('create_interview'))
    # Giriş yapmamış kullanıcıyı login sayfasına yönlendir
    return redirect(url_for('login'))

@app.route('/join')
def join():
    return render_template('interview_entry.html')

@app.route('/interview')
def interview():
    code = request.args.get('code')
    if not code:
        return render_template('error.html', error="Mülakat kodu gereklidir"), 400
        
    try:
        # JSON dosyasından mülakat verilerini al
        json_path = os.path.join('interviews', f'{code}.json')
        if not os.path.exists(json_path):
            return render_template('error.html', error="Geçersiz mülakat kodu"), 404
            
        with open(json_path, 'r', encoding='utf-8') as f:
            interview_data = json.load(f)
            
        # Mülakat durumunu kontrol et
        if interview_data.get('status') == 'completed' or interview_data.get('ended'):
            return render_template('error.html', 
                error="Bu mülakat daha önce tamamlanmış. Aynı mülakat kodunu tekrar kullanamazsınız."), 403
            
        # Interview data'yı güncelle
        interview_data['status'] = 'active'
        interview_data['started_at'] = datetime.now().isoformat()
        
        # Değişiklikleri kaydet
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(interview_data, f, ensure_ascii=False, indent=2)
            
        # Global current_interview'ı oluştur
        global current_interview
        current_interview = InterviewAssistant()
        current_interview.set_interview_details(code)
        
        # Adayın adını CV'den al veya varsayılan değer kullan
        candidate_name = interview_data.get("candidate_name", "İsimsiz Aday")
        
        # CV yüklendiyse ve isim CV'den alınabilirse, öncelikle bunu kullan
        if interview_data.get("cv_data") and interview_data["cv_data"].get("personal_info"):
            cv_name = interview_data["cv_data"]["personal_info"].get("name")
            if cv_name and cv_name.strip():
                candidate_name = cv_name
        
        # Log ile aday adını kontrol et
        logger.info(f"CV verileri başarıyla yüklendi - Mülakat Kodu: {code}")
        logger.info(f"Mülakat detayları ayarlandı: {code}")
        
        return render_template('interview.html', 
                             interview={
                             "candidate_name": candidate_name,
                             "position": interview_data.get("position", "Belirtilmemiş Pozisyon"),
                                 "code": code,
                             "created_at": interview_data.get("created_at", "")
                             })
                             
    except Exception as e:
        logger.error(f"Mülakat sayfası yükleme hatası: {str(e)}")
        return render_template('error.html', error="Mülakat yüklenirken bir hata oluştu"), 500

@app.route('/create_interview', methods=['GET', 'POST'])
@login_required
def create_interview():
    """Mülakat oluştur sayfası"""
    if request.method == 'GET':
        return render_template('create_interview.html')
        
    try:
        data = request.get_json()
        candidate_name = data.get('candidate_name')
        position = data.get('position')
        max_questions = data.get('max_questions', 5)  # Varsayılan soru sayısı 5
        cv_data = data.get('cv_data', {})  # CV verilerini al
        
        if not candidate_name or not position:
            return jsonify({
                'success': False,
                'error': 'Aday adı ve pozisyon gereklidir'
            })
            
        # Kullanıcının özel soruları varsa onları kullan
        custom_questions = data.get('custom_questions')
        
        if not custom_questions:
            try:
                # CV verisi varsa bunu prompt'a ekle
                cv_context = ""
                if cv_data:
                    cv_context = f"""
                    Aday CV Bilgileri:
                    - Beceriler: {', '.join(cv_data.get('skills', []))}
                    - Deneyim: {', '.join([f"{exp.get('title')} - {exp.get('company')}" for exp in cv_data.get('experience', [])])}
                    - Eğitim: {', '.join([f"{edu.get('degree')} - {edu.get('institution')}" for edu in cv_data.get('education', [])])}
                    
                    CV bilgilerini dikkate alarak adayın deneyimlerine ve becerilerine uygun, derinlemesine sorular sor.
                    """
                
                # GPT ile pozisyona özel sorular oluştur
                prompt = f"""
                {position} pozisyonu için {max_questions} adet mülakat sorusu oluştur.
                
                Aday Bilgileri:
                - İsim: {candidate_name}
                - Pozisyon: {position}
                {cv_context}
                
                Önemli Kurallar:
                1. Kesinlikle "1. soru", "2. soru" gibi numaralandırma kullanma
                2. Her soru doğal bir sohbet akışının parçası olmalı
                3. Sorular şu konuları kapsamalı:
                   - Genel deneyim ve motivasyon
                   - Pozisyona özel teknik bilgi
                   - Problem çözme yaklaşımı
                   - Takım çalışması ve iletişim
                
                Örnek Doğal Soru Formatı:
                - "Bize biraz kendinizden ve kariyerinizden bahseder misiniz?"
                - "Bu pozisyona başvurmanızın arkasındaki motivasyonunuz nedir?"
                - "[Teknik konu] hakkındaki deneyimlerinizi paylaşır mısınız?"
                
                Lütfen tam olarak {max_questions} adet doğal ve akıcı soru oluştur. Sorular Türkçe olmalı ve bir sohbet akışı içinde sorulabilecek şekilde olmalı."""
                
                response = openai_client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "Sen bir kıdemli İK uzmanısın. Doğal ve etkili mülakat soruları oluşturuyorsun."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=1000
                )
                
                # GPT yanıtından soruları al
                custom_questions = response.choices[0].message.content.strip().split('\n')
                # Boş satırları temizle
                custom_questions = [q.strip() for q in custom_questions if q.strip()]
                
                # Soru sayısını sınırla (eğer GPT daha fazla soru döndüyse)
                if len(custom_questions) > max_questions:
                    custom_questions = custom_questions[:max_questions]
                
                logger.info(f"GPT tarafından {len(custom_questions)} soru oluşturuldu")
                
            except Exception as e:
                logger.error(f"GPT soru oluşturma hatası: {str(e)}")
                # Hata durumunda varsayılan soruları kullan
                default_questions = [
                    "Bize biraz kendinizden ve kariyerinizden bahseder misiniz?",
                    "Bu pozisyona başvurmanızın arkasındaki motivasyonunuz nedir?",
                    "Şimdiye kadar karşılaştığınız en zorlu teknik problemi ve nasıl çözdüğünüzü anlatır mısınız?",
                    "Takım çalışması deneyimlerinizden bahseder misiniz?",
                    "Bize gelecekteki kariyer hedeflerinizden bahseder misiniz?",
                    "Bu pozisyonun gereksinimlerine ne kadar uygun olduğunuzu düşünüyorsunuz?",
                    "Önceki rollerinizde hangi başarıları elde ettiniz?",
                    "Bu sektörde karşılaşılan zorluklar hakkında ne düşünüyorsunuz?"
                ]
                # Varsayılan soruları soru sayısına göre sınırla
                custom_questions = default_questions[:max_questions]
        else:
            logger.info(f"Kullanıcı tarafından {len(custom_questions)} soru sağlandı")
            
        # Benzersiz kod oluştur
        interview_code = generate_interview_code()
        
        # Mülakat verilerini hazırla
        interview_info = {
            "candidate_info": {
                "name": candidate_name,
                "position": position,
                "requirements": data.get('requirements', []),
                "custom_questions": custom_questions
            },
            "created_at": datetime.now().isoformat(),
            "expires_at": (datetime.now() + timedelta(hours=24)).isoformat(),
            "status": "active"
        }
        
        try:
            # Memory'ye kaydet
            interview_data[interview_code] = interview_info
            
            # JSON dosyasına kaydet
            json_path = os.path.join('interviews', f'{interview_code}.json')
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump({
                    "code": interview_code,
                    "candidate_name": candidate_name,
                    "position": position,
                    "requirements": interview_info["candidate_info"]["requirements"],
                    "questions": custom_questions,
                    "created_at": interview_info["created_at"],
                    "status": "active"
                }, f, ensure_ascii=False, indent=2)
                
            logger.info(f"Yeni mülakat oluşturuldu: {interview_code}")
            
            # Mülakat linkini oluştur
            domain = os.getenv('DOMAIN_NAME', request.host_url.rstrip('/'))
            
            # HTTPS kontrolü
            protocol = 'https' if request.is_secure or 'https' in request.host_url else 'http'
            
            # Ana domain ve alt domainler için farklı URL'ler oluştur
            interview_urls = {
                'main': f"{protocol}://{domain}/interview?code={interview_code}",
                'create': f"{protocol}://{domain}/",
                'join': f"{protocol}://{domain}/join"
            }
            
            return jsonify({
                'success': True,
                'code': interview_code,
                'urls': interview_urls,
                'questions': custom_questions
            })
            
        except Exception as e:
            logger.error(f"Mülakat kaydetme hatası: {str(e)}")
            return jsonify({
                'success': False,
                'error': 'Mülakat kaydedilemedi'
            })
            
    except Exception as e:
        logger.error(f"Mülakat oluşturma hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/verify_code', methods=['POST'])
def verify_code():
    """Mülakat kodunu doğrula"""
    try:
        data = request.get_json()
        code = data.get('code')

        if not code:
            return jsonify({
                'success': False,
                'error': 'Mülakat kodu gerekli'
            })

        try:
            # Önce memory'de kontrol et
            if code in interview_data:
                return jsonify({'success': True})
                
            # JSON dosyasından mülakat kodu kontrolü
            json_path = os.path.join('interviews', f'{code}.json')
            if os.path.exists(json_path):
                with open(json_path, 'r', encoding='utf-8') as f:
                    interview_json = json.load(f)
                    
                # Interview data'yı güncelle
                interview_data[code] = {
                    "candidate_info": {
                        "name": interview_json.get("candidate_name"),
                        "position": interview_json.get("position"),
                        "requirements": interview_json.get("requirements", []),
                        "custom_questions": interview_json.get("questions", [])
                    },
                    "created_at": interview_json.get("created_at", datetime.now().isoformat()),
                    "expires_at": (datetime.now() + timedelta(hours=24)).isoformat(),
                    "status": "active"
                }
                
                return jsonify({'success': True})
            
            # Quiz kodu kontrolü
            for file in os.listdir(QUESTIONS_DIR):
                if file.endswith(f"_{code}.json"):
                    return jsonify({'success': True})
                
            return jsonify({
                'success': False,
                'error': 'Geçersiz mülakat veya quiz kodu'
            })
            
        except Exception as e:
            logger.error(f"Kod doğrulama işlem hatası: {str(e)}")
            return jsonify({
                'success': False,
                'error': 'Kod doğrulanamadı'
            })
            
    except Exception as e:
        logger.error(f"Kod doğrulama genel hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/check_quiz_code', methods=['GET'])
def check_quiz_code():
    """Verilen kodun quiz kodu olup olmadığını kontrol et"""
    try:
        code = request.args.get('code', '')
        if not code:
            return jsonify({
                'success': False,
                'error': 'Kod gerekli',
                'is_quiz': False
            })
        
        # Quiz klasöründe bu koduyla biten dosya ara
        for file in os.listdir(QUESTIONS_DIR):
            if file.endswith(f"_{code}.json"):
                return jsonify({
                    'success': True,
                    'is_quiz': True
                })
                
        # Eğer bulunmadıysa, mülakat kodu olduğunu varsay
        return jsonify({
            'success': True,
            'is_quiz': False
        })
        
    except Exception as e:
        logger.error(f"Quiz kodu kontrol hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'is_quiz': False
        })

@app.route('/start_interview', methods=['POST'])
async def start_interview():
    try:
        data = request.json
        if not data or 'code' not in data:
            return jsonify({
                "success": False,
                "error": "Mülakat kodu gerekli"
            }), 400

        global current_interview
        current_interview = InterviewAssistant()
        current_interview.set_interview_details(data['code'])
        
        # Hoşgeldin mesajını al
        welcome_message = "Merhaba, mülakata başlamak için hazır olduğunuzda 'Başlayalım' diyebilirsiniz."
        
        logger.info(f"Mülakat başlatıldı: {data['code']}")
        
        return jsonify({
            "success": True,
            "welcome_message": welcome_message
        })
        
    except Exception as e:
        logger.error(f"Mülakat başlatma hatası: {str(e)}")
        return jsonify({
            "success": False,
            "error": f"Mülakat başlatılamadı: {str(e)}"
        }), 500

@app.route('/start_recording', methods=['POST'])
async def start_recording():
    try:
        if not current_interview:
            return jsonify({
                "success": False,
                "error": "Lütfen önce mülakatı başlatın"
            }), 400
            
        # Ses kaydını başlat
        recording = current_interview.record_audio()
        if recording:
            filename = current_interview.save_audio(recording)
            return jsonify({
                "success": True,
                "message": "Ses kaydı başarıyla tamamlandı",
                "filename": filename
            })
        else:
            return jsonify({
                "success": False,
                "error": "Ses kaydı alınamadı"
            }), 400
        
    except Exception as e:
        logger.error(f"Ses kaydı başlatma hatası: {str(e)}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route('/process_audio', methods=['POST'])
async def process_audio():
    try:
        audio_file = request.files.get('audio')
        if not audio_file:
            return jsonify({
                'success': False,
                'error': 'Ses dosyası gerekli'
            })

        # Geçici dizini kontrol et
        if not os.path.exists('temp'):
            os.makedirs('temp')

        # Ses dosyasını kaydet
        timestamp = int(time.time())
        temp_path = f'temp/audio_{timestamp}_{random.randint(100000, 999999)}.webm'
        audio_file.save(temp_path)
        logger.info(f"WebM dosyası kaydedildi: {os.path.abspath(temp_path)}")

        try:
            # Ses dosyasını metne çevir
            transcript = await current_interview.transcribe_audio(temp_path)
            if not transcript:
                logger.warning("Ses tanıma başarısız oldu")
                return jsonify({
                    'success': False,
                    'error': 'Ses tanınamadı',
                    'continue_listening': True  # Dinlemeye devam et
                })

            logger.info(f"Tanınan metin: {transcript}")

            # GPT yanıtını al
            gpt_response, is_interview_ended = await current_interview.get_gpt_response(transcript)
            if not gpt_response:
                logger.warning("GPT yanıtı alınamadı")
                return jsonify({
                    'success': False,
                    'error': 'GPT yanıtı alınamadı',
                    'continue_listening': True  # Dinlemeye devam et
                })

            logger.info(f"GPT yanıtı: {gpt_response}")
            
            # Mülakat bittiğinde istemciye bildir
            interview_ended = is_interview_ended or "MÜLAKAT_BİTTİ" in gpt_response
            
            # Eğer mülakat bittiyse, "MÜLAKAT_BİTTİ" cümlesini yanıttan çıkar
            if "MÜLAKAT_BİTTİ" in gpt_response:
                gpt_response = gpt_response.replace("MÜLAKAT_BİTTİ", "").strip()
            
            # İstemciye yanıt gönder
            return jsonify({
                'success': True,
                'transcript': transcript,
                'response': gpt_response,
                'interview_ended': interview_ended
            })
            
        except Exception as e:
            logger.error(f"Ses işleme hatası: {str(e)}")
            return jsonify({
                'success': False,
                'error': str(e),
                'continue_listening': True  # Genel hata durumunda bile dinlemeye devam et
            })

        finally:
            # Geçici dosyayı temizle
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                    logger.info("Geçici ses dosyası silindi")
                except Exception as e:
                    logger.warning(f"Geçici dosya silme hatası: {str(e)}")

    except Exception as e:
        logger.error(f"Genel hata: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'continue_listening': True  # Genel hata durumunda bile dinlemeye devam et
        })


@app.route('/generate_report', methods=['POST'])
async def generate_report():
    try:
        logging.info("Rapor oluşturma süreci başlatıldı...")
        
        # İstek verilerini al
        data = request.get_json()
        interview_code = data.get('interview_code')
        conversation_history = data.get('conversation_history', [])
        
        if not interview_code:
            return jsonify({'success': False, 'error': 'Mülakat kodu gerekli'}), 400
            
        # Mülakat dosyasını kontrol et
        interview_file = f'interviews/{interview_code}.json'
        if not os.path.exists(interview_file):
            return jsonify({'success': False, 'error': 'Mülakat bulunamadı'}), 404
            
        # Mülakat verilerini oku
        with open(interview_file, 'r', encoding='utf-8') as f:
            interview_data = json.load(f)
            
        if not conversation_history:
            return jsonify({'success': False, 'error': 'Konuşma geçmişi boş'}), 400
            
        # GPT değerlendirmesi için prompt hazırla
        prompt = f"""
        Aşağıdaki mülakat için detaylı bir değerlendirme raporu hazırla:
        
        Aday: {interview_data.get('candidate_name')}
        Pozisyon: {interview_data.get('position')}
        Tarih: {interview_data.get('created_at')}
        
        Mülakat Konuşma Akışı:
        """
        
        # Konuşma akışını ekle
        for message in conversation_history:
            role = "Aday" if message['role'] == 'user' else "Mülakat Asistanı"
            prompt += f"\n{role}: {message['content']}\n"
            
        prompt += """
        Lütfen aşağıdaki kriterlere göre değerlendirme yap:
        1. Teknik Yetkinlik (1-10)
        2. İletişim Becerileri (1-10)
        3. Problem Çözme (1-10)
        4. Deneyim Seviyesi (1-10)
        
        Her kriter için detaylı açıklama ve örnekler ver.
        Güçlü yönler ve gelişim alanlarını belirt.
        Genel değerlendirme ve tavsiyeler ekle.
        """
        
        # GPT'den değerlendirme al
        client = OpenAI()
        response = client.chat.completions.create(
            model="gpt-4-turbo-preview",
            messages=[{"role": "system", "content": "Sen bir mülakat değerlendirme uzmanısın."},
                     {"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=2000
        )
        
        evaluation_text = response.choices[0].message.content
        
        # Skorları ve bölümleri çıkar
        scores = {
            'technical': extract_score(evaluation_text, 'Teknik Yetkinlik'),
            'communication': extract_score(evaluation_text, 'İletişim Becerileri'),
            'problem_solving': extract_score(evaluation_text, 'Problem Çözme'),
            'experience': extract_score(evaluation_text, 'Deneyim Seviyesi')
        }
        
        sections = {
            'strengths': extract_section(evaluation_text, 'Güçlü Yönler'),
            'improvements': extract_section(evaluation_text, 'Gelişim Alanları'),
            'overall': extract_section(evaluation_text, 'Genel Değerlendirme')
        }
        
        # Genel skoru hesapla
        overall_score = calculate_overall_score(scores)
        
        # PDF rapor dosya adını oluştur
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        pdf_filename = f"report_{interview_code}_{timestamp}.pdf"
        pdf_path = os.path.join('reports', pdf_filename)
        
        # PDF raporu oluştur
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        
        # Başlık
        elements.append(Paragraph(f"Mülakat Değerlendirme Raporu", styles['Title']))
        elements.append(Spacer(1, 12))
        
        # Aday Bilgileri
        elements.append(Paragraph(f"Aday: {interview_data.get('candidate_name')}", styles['Heading2']))
        elements.append(Paragraph(f"Pozisyon: {interview_data.get('position')}", styles['Normal']))
        elements.append(Paragraph(f"Tarih: {interview_data.get('created_at')}", styles['Normal']))
        elements.append(Spacer(1, 12))
        
        # Skorlar
        elements.append(Paragraph("Değerlendirme Skorları", styles['Heading2']))
        for category, score in scores.items():
            elements.append(Paragraph(f"{category.replace('_', ' ').title()}: {score}/10", styles['Normal']))
        elements.append(Paragraph(f"Genel Skor: {overall_score}/10", styles['Normal']))
        elements.append(Spacer(1, 12))
        
        # Detaylı Değerlendirme
        elements.append(Paragraph("Detaylı Değerlendirme", styles['Heading2']))
        for section, content in sections.items():
            elements.append(Paragraph(f"{section.replace('_', ' ').title()}", styles['Heading3']))
            elements.append(Paragraph(content, styles['Normal']))
            elements.append(Spacer(1, 12))
            
        # Konuşma Akışı
        elements.append(Paragraph("Mülakat Konuşma Akışı", styles['Heading2']))
        for message in conversation_history:
            role = "Aday" if message['role'] == 'user' else "Mülakat Asistanı"
            elements.append(Paragraph(f"{role}:", styles['Heading4']))
            elements.append(Paragraph(message['content'], styles['Normal']))
            elements.append(Spacer(1, 6))
        
        # PDF'i oluştur
        doc.build(elements)
        
        # Webhook URL'sini kontrol et
        webhook_url = interview_data.get('webhook_rapor_url') or os.getenv('WEBHOOK_RAPOR_URL')
        
        if webhook_url:
            # Webhook verilerini hazırla
            webhook_data = {
                'interview_code': interview_code,
                'candidate_name': interview_data.get('candidate_name'),
                'position': interview_data.get('position'),
                'date': interview_data.get('created_at'),
                'evaluation': {
                    'scores': scores,
                    'overall_score': overall_score,
                    'sections': sections
                },
                'pdf_path': pdf_path
            }
            
            try:
                # Webhook'u gönder
                webhook_response = requests.post(webhook_url, json=webhook_data, timeout=5)
                webhook_response.raise_for_status()
                logging.info(f"Webhook başarıyla gönderildi: {webhook_response.status_code}")
            except Exception as e:
                logging.error(f"Webhook gönderimi başarısız: {str(e)}")
        
        return jsonify({
            'success': True,
            'pdf_path': pdf_path,
            'evaluation': {
                'scores': scores,
                'overall_score': overall_score,
                'sections': sections
            }
        })
        
    except Exception as e:
        logging.error(f"Rapor oluşturma hatası: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

def extract_score(text, category):
    """Metinden puan çıkarma"""
    try:
        pattern = f"{category}.*?(\\d+)"
        match = re.search(pattern, text, re.DOTALL)
        if match:
            score = int(match.group(1))
            return min(max(score, 0), 100)  # 0-100 arası sınırla
    except:
        pass
    return 0

def extract_section(text, section):
    """Metinden bölüm çıkarma"""
    try:
        pattern = f"{section}:(.*?)(?=\n\n|$)"
        match = re.search(pattern, text, re.DOTALL)
        if match:
            return match.group(1).strip()
    except:
        pass
    return ""

def calculate_overall_score(evaluation_data):
    """Genel puanı hesapla"""
    weights = {
        "teknik_bilgi_ve_deneyim": 0.4,
        "iletisim_becerileri": 0.2,
        "problem_cozme": 0.2,
        "ozguven_profesyonellik": 0.2
    }
    
    total_score = 0
    for key, weight in weights.items():
        total_score += evaluation_data.get(key, 0) * weight
        
    return round(total_score, 1)

@app.route('/check_audio_support', methods=['GET'])
def check_audio_support():
    """Desteklenen ses formatlarını kontrol et"""
    try:
        supported_formats = {
            'webm': True,
            'wav': True,
            'mp3': True,
            'ogg': True
        }
        return jsonify({
            'success': True,
            'supported_formats': supported_formats
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# Veritabanı bağlantısı için yardımcı fonksiyon
def get_db_connection():
    try:
        conn = sqlite3.connect('data/interview.db')
        conn.row_factory = sqlite3.Row
        return conn
    except Exception as e:
        logger.error(f"Veritabanı bağlantı hatası: {str(e)}")
        raise

def create_or_get_interview_code(email):
    """Email adresine göre mülakat kodu oluştur veya var olanı getir"""
    conn = get_db_connection()
    try:
        # Önce mevcut kodu kontrol et
        cursor = conn.execute('SELECT code FROM interview_codes WHERE email = ?', (email,))
        result = cursor.fetchone()
        
        if result:
            return result['code']
            
        # Yeni kod oluştur
        while True:
            # Email'den benzersiz bir kod oluştur
            hash_object = hashlib.md5(email.encode())
            code = hash_object.hexdigest()[:6].upper()
            
            # Kodun benzersiz olduğunu kontrol et
            cursor = conn.execute('SELECT code FROM interview_codes WHERE code = ?', (code,))
            if not cursor.fetchone():
                break
        
        # Yeni kodu kaydet
        conn.execute(
            'INSERT INTO interview_codes (email, code, created_at) VALUES (?, ?, ?)',
            (email, code, datetime.now().isoformat())
        )
        conn.commit()
        return code
    finally:
        conn.close()

def save_interview_data(data, code):
    """Mülakat verilerini JSON dosyası olarak kaydet"""
    try:
        # Veriyi yeni formata dönüştür
        formatted_data = {
            "code": code,
            "candidate_name": data.get('adSoyad') or data.get('candidate_name'),
            "position": data.get('isIlaniPozisyonu') or data.get('position'),
            "questions": data.get('mulakatSorulari') or data.get('questions', [
                "1. Yapay zekanın temel bileşenleri hakkında bilgi verebilir misiniz?",
                "2. Belirli bir veri setinde overfitting problemini nasıl tanımlar ve çözersiniz?",
                "3. Günlük çalışmalarınızda hangi yapay zeka frameworklerini kullandınız?",
                "4. Çeşitli regresyon teknikleri hakkında bilgi verebilir misiniz?",
                "5. NLP konusunda ne gibi deneyimleriniz var?"
            ]),
            "status": data.get('status', 'pending'),
            "created_at": data.get('created_at', datetime.now().isoformat()),
            "cv_data": data.get('cv_data', {})  # CV verilerini ekle
        }
        
        file_path = os.path.join('interviews', f'{code}.json')
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(formatted_data, f, ensure_ascii=False, indent=2)
            
        logger.info(f"Mülakat verisi kaydedildi: {code}")
        return file_path
    except Exception as e:
        logger.error(f"JSON kaydetme hatası: {str(e)}")
        return None

def send_webhook_notification(webhook_url, data):
    """Webhook'a bildirim gönder"""
    try:
        response = requests.post(
            webhook_url,
            json=data,
            headers={'Content-Type': 'application/json'}
        )
        return response.status_code == 200
    except Exception as e:
        logger.error(f"Webhook gönderme hatası: {str(e)}")
        return False

@app.route('/webhook/interview', methods=['POST'])
def webhook_aday_handler():
    try:
        data = request.json
        if not data:
            return jsonify({"error": "Veri bulunamadı"}), 400

        # Gerekli alanları kontrol et
        required_fields = ['adSoyad', 'isIlaniPozisyonu']
        for field in required_fields:
            if field not in data:
                return jsonify({"error": f"Eksik alan: {field}"}), 400

        # Mülakat kodu oluştur
        interview_code = generate_interview_code()
        
        # Webhook URL'lerini al
        webhook_aday_url = data.get('webhook_aday_url', WEBHOOK_ADAY_URL)
        webhook_rapor_url = data.get('webhook_rapor_url', WEBHOOK_RAPOR_URL)
        
        # JSON dosyasını oluştur
        interview_data = {
            "code": interview_code,
            "candidate_name": data.get("adSoyad"),
            "position": data.get("isIlaniPozisyonu"),
            "requirements": data.get("isIlaniGereksinimleri", []),
            "questions": data.get("mulakatSorulari", []),
            "created_at": datetime.now().isoformat(),
            "status": "pending",
            "webhook_aday_url": webhook_aday_url,
            "webhook_rapor_url": webhook_rapor_url
        }
        
        # JSON dosyasını kaydet
        json_path = os.path.join('interviews', f'{interview_code}.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(interview_data, f, ensure_ascii=False, indent=2)
        
        # WebhookAdayİşlem'e yanıt gönder
        response_data = {
            "success": True,
            "interview_code": interview_code,
            "interview_url": f"{request.host_url}interview?code={interview_code}",
            "expires_at": (datetime.now() + timedelta(hours=24)).isoformat()
        }
        
        # Webhook'a bildirim gönder
        try:
            if webhook_aday_url:
                webhook_response = requests.post(
                    webhook_aday_url,
                    json=response_data,
                    headers={'Content-Type': 'application/json'},
                    timeout=5  # 5 saniye timeout
                )
                logger.info(f"WebhookAdayİşlem yanıtı: {webhook_response.status_code}")
        except requests.exceptions.RequestException as e:
            logger.error(f"WebhookAdayİşlem gönderme hatası: {str(e)}")
            # Webhook hatası olsa bile işleme devam et

        return jsonify(response_data)

    except Exception as e:
        logger.error(f"Webhook alım hatası: {str(e)}")
        return jsonify({"error": str(e)}), 500

def send_report_webhook(pdf_path, evaluation_data):
    """Raporu webhook'a gönder"""
    try:
        # Mülakat verilerini hazırla
        interview_summary = {
            "mulakat_kodu": current_interview.code,
            "aday_bilgileri": {
                "isim": current_interview.candidate_name,
                "pozisyon": current_interview.position,
                "cv_ozeti": current_interview.cv_summary,
                "on_bilgi": current_interview.pre_info,
                "gereksinimler": current_interview.requirements,
                "tarih": current_interview.start_time.isoformat()
            },
            "mulakat_metrikleri": {
                "iletisim_puani": current_interview.metrics["iletisim_puani"],
                "ozguven_puani": current_interview.metrics["ozguven_puani"],
                "teknik_bilgi": current_interview.metrics["teknik_bilgi"],
                "genel_puan": current_interview.metrics["genel_puan"]
            },
            "degerlendirme": evaluation_data,
            "konusma_akisi": current_interview.conversation_history,
            "rapor_dosyasi": pdf_path,
            "olusturulma_tarihi": datetime.now().isoformat()
        }
        
        # JSON dosyasından webhook URL'sini kontrol et
        json_path = os.path.join('interviews', f'{current_interview.code}.json')
        webhook_url = WEBHOOK_RAPOR_URL  # Varsayılan URL
        
        if os.path.exists(json_path):
            with open(json_path, 'r', encoding='utf-8') as f:
                interview_json = json.load(f)
                if 'webhook_rapor_url' in interview_json:
                    webhook_url = interview_json['webhook_rapor_url']
                    logger.info(f"JSON'dan webhook URL'si kullanılıyor: {webhook_url}")
                else:
                    logger.info(f"Varsayılan webhook URL'si kullanılıyor: {webhook_url}")
        
        # WebhookRapor'a gönder
        response = requests.post(
            webhook_url,
            json=interview_summary,
            headers={
                'Content-Type': 'application/json',
                'X-Interview-Code': current_interview.code
            }
        )
        
        if response.status_code == 200:
            logger.info(f"Rapor webhook'a başarıyla gönderildi: {current_interview.code}")
            return True
        else:
            logger.error(f"Webhook hatası: {response.status_code} - {response.text}")
            return False
                
    except Exception as e:
        logger.error(f"Webhook gönderme hatası: {str(e)}")
        return False

@app.route('/reports/<path:filename>')
def serve_report(filename):
    """Mülakat raporlarını sunmak için endpoint"""
    try:
        return send_from_directory('reports', filename)
    except Exception as e:
        logging.error(f"Rapor sunma hatası: {str(e)}")
        return "Rapor bulunamadı veya erişim hatası", 404

@app.route('/analyze_cv', methods=['POST'])
def analyze_cv():
    """CV dosyasını analiz eden endpoint"""
    try:
        if 'cv_file' not in request.files:
            return jsonify({
                'success': False,
                'error': 'Dosya bulunamadı'
            })
        
        cv_file = request.files['cv_file']
        position = request.form.get('position', '')
        
        # Eğer pozisyon belirtilmemişse kullanıcıya sormak için bilgi gönder
        if not position:
            return jsonify({
                'success': True,
                'needs_position': True,
                'message': 'Hangi pozisyon için CV değerlendirmesi yapılacağını belirtiniz.'
            })
        
        if cv_file.filename == '':
            return jsonify({
                'success': False,
                'error': 'Dosya seçilmedi'
            })
            
        # Geçici dosyayı kaydet
        temp_path = os.path.join('temp', f'cv_{int(time.time())}_{cv_file.filename}')
        os.makedirs('temp', exist_ok=True)
        cv_file.save(temp_path)
        
        # Dosya tipini algıla
        file_ext = os.path.splitext(cv_file.filename)[1].lower()
        
        # Dosya içeriğini çıkar
        cv_text = ""
        if file_ext == '.pdf':
            # PDF işleme kodu - PyPDF2 kullanılabilir
            import PyPDF2
            try:
                with open(temp_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    cv_text = ""
                    for page in reader.pages:
                        cv_text += page.extract_text() + "\n"
            except Exception as e:
                logging.error(f"PDF işleme hatası: {str(e)}")
                cv_text = "PDF işlenirken bir hata oluştu."
                
        elif file_ext in ['.docx', '.doc']:
            # DOCX işleme kodu - python-docx kullanılabilir
            import docx
            try:
                doc = docx.Document(temp_path)
                cv_text = ""
                for para in doc.paragraphs:
                    cv_text += para.text + "\n"
            except Exception as e:
                logging.error(f"DOCX işleme hatası: {str(e)}")
                cv_text = "DOCX işlenirken bir hata oluştu."
        else:
            # Desteklenmeyen dosya formatı
            return jsonify({
                'success': False,
                'error': 'Desteklenmeyen dosya formatı'
            })
            
        # OpenAI ile CV analizi yap
        client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        
        # Pozisyon hakkında daha fazla bilgi soruyorsa, pozisyon gereksinimleri analizi yap
        position_context = ""
        if position:
            try:
                position_response = client.chat.completions.create(
                    model="gpt-4-turbo",
                    response_format={"type": "json_object"},
                    messages=[
                        {"role": "system", "content": """
                            Verilen pozisyonun gerekenliklerini analiz eden bir İK uzmanısın. 
                            Pozisyon adından yola çıkarak aşağıdaki formatta JSON çıktısı üreteceksin:
                            
                            {
                                "position_title": "Pozisyon adı",
                                "required_skills": ["beceri1", "beceri2", ...],
                                "preferred_experience": ["deneyim1", "deneyim2", ...],
                                "education_requirements": ["eğitim1", "eğitim2", ...],
                                "keywords": ["anahtar kelime1", "anahtar kelime2", ...]
                            }
                            
                            Tahminlerin mümkün olduğunca gerçekçi ve doğru olmalı.
                        """},
                        {"role": "user", "content": f"Bu pozisyon için gereksinimleri analiz et: {position}"}
                    ]
                )
                
                position_data = json.loads(position_response.choices[0].message.content)
                position_context = f"""
                    Pozisyon bilgilerini de analiz et:
                    Pozisyon: {position}
                    Gerekli beceriler: {', '.join(position_data.get('required_skills', []))}
                    Tercih edilen deneyimler: {', '.join(position_data.get('preferred_experience', []))}
                    Anahtar kelimeler: {', '.join(position_data.get('keywords', []))}
                """
            except Exception as e:
                logging.error(f"Pozisyon analizi hatası: {str(e)}")
        
        # CV Analizi yap
        try:
            response = client.chat.completions.create(
                model="gpt-4-turbo",
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": f"""
                    Sen bir CV analiz uzmanısın. Verilen CV metnini analiz ederek aşağıdaki yapılandırılmış JSON formatında döndür:
                    
                    {{
                        "personal_info": {{
                            "name": "Ad Soyad",
                            "email": "e-posta adresi",
                            "phone": "telefon numarası",
                            "location": "konum"
                        }},
                        "summary": "Adayın kısa özeti",
                        "skills": ["beceri1", "beceri2", ...],
                        "experience": [
                            {{
                                "title": "Pozisyon başlığı",
                                "company": "Şirket adı",
                                "duration": "Çalışma süresi",
                                "description": "Pozisyon açıklaması",
                                "achievements": ["başarı1", "başarı2", ...]
                            }},
                            ...
                        ],
                        "education": [
                            {{
                                "degree": "Derece",
                                "institution": "Okul/Üniversite adı",
                                "year": "Eğitim yılı",
                                "description": "Eğitim açıklaması"
                            }},
                            ...
                        ],
                        "languages": ["dil1", "dil2", ...],
                        "certificates": ["sertifika1", "sertifika2", ...],
                        "profile_summary": "Bu adayın güçlü yanları ve pozisyona uygunluğu hakkında 3-4 cümlelik özet",
                        "position_match": {{
                            "match_percentage": 0-100 arası bir sayı,
                            "matching_skills": ["eşleşen beceri1", ...],
                            "missing_skills": ["eksik beceri1", ...],
                            "recommendations": ["öneri1", "öneri2", ...]
                            }},
                            "strengths": ["güçlü yön1", "güçlü yön2", ...],
                            "areas_to_improve": ["geliştirilecek alan1", "geliştirilecek alan2", ...],
                            "suggested_questions": ["soru1", "soru2", "soru3", "soru4", "soru5"]
                    }}
                    
                    Bazı alanlar CV'de bulunmayabilir. Bu durumda boş dizi veya null değeri kullan.
                    CV'deki bilgileri mümkün olduğunca eksiksiz çıkar. Hiçbir bilgi uydurmadan, CV'de olan bilgileri yapılandırılmış formata çevir.
                        Puanlamayı çok titiz ve gerçekçi yap, abartılı pozitif değerlendirmelerden kaçın.
                        Güçlü yönler, geliştirilecek alanlar ve önerilen soruları da mutlaka ekle.
                        
                    {position_context}
                """},
                {"role": "user", "content": f"Aşağıdaki CV'yi analiz et ve yapılandırılmış JSON formatında döndür:\n\n{cv_text}"}
            ]
            )
            
            # JSON yanıtını parse et
            cv_data = json.loads(response.choices[0].message.content)
        except Exception as e:
            logging.error(f"CV analiz API hatası: {str(e)}")
            return jsonify({
                'success': False,
                'error': f"CV analiz edilirken bir API hatası oluştu: {str(e)}"
            })
        
        # Geçici dosyayı sil
        try:
            os.remove(temp_path)
        except Exception as e:
            logging.error(f"Geçici dosya silme hatası: {str(e)}")
            
        # CV'ye özel sorular oluştur
        try:
            position_title = position or (cv_data.get('experience') and cv_data.get('experience')[0].get('title', ''))
            skills = cv_data.get('skills', [])
            
            # CV'ye özel mülakat soruları oluştur
            questions_prompt = f"""
                Bu adayın CV'sine göre en etkili mülakat sorularını oluştur:
                
                Aday Bilgileri:
                - İsim: {cv_data.get('personal_info', {}).get('name', 'Aday')}
                - Beceriler: {', '.join(skills[:5])}
                - Deneyim: {', '.join([f"{exp.get('title')} - {exp.get('company')}" for exp in cv_data.get('experience', [])[:2]])}
                - Eğitim: {', '.join([f"{edu.get('degree')} - {edu.get('institution')}" for edu in cv_data.get('education', [])[:1]])}
                
                Pozisyon: {position_title}
                
                Aday CV'sine göre özelleştirilmiş, derinlemesine 6-8 mülakat sorusu oluştur. Sorular numaralandırılmasın ve her biri bir paragraf olacak şekilde döndürülsün.
            """
            
            questions_response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Sen bir kıdemli İK uzmanısın ve işe alım mülakatları için soru hazırlıyorsun."},
                    {"role": "user", "content": questions_prompt}
                ],
                temperature=0.7,
                max_tokens=1000
            )
            
            # GPT tarafından oluşturulan soruları al
            generated_questions = questions_response.choices[0].message.content.strip().split('\n')
            generated_questions = [q.strip() for q in generated_questions if q.strip() and len(q.strip()) > 10]
            
            # Soruları CV verilerine ekle
            cv_data['generated_questions'] = generated_questions
            
        except Exception as e:
            logging.error(f"Soru oluşturma hatası: {str(e)}")
            cv_data['generated_questions'] = []
            
        # Başarılı analiz sonucunu döndür
        result = {
            'success': True,
            'personal_info': cv_data.get('personal_info', {}),
            'summary': cv_data.get('summary', ''),
            'skills': cv_data.get('skills', []),
            'experience': cv_data.get('experience', []),
            'education': cv_data.get('education', []),
            'languages': cv_data.get('languages', []),
            'certificates': cv_data.get('certificates', []),
            'profile_summary': cv_data.get('profile_summary', ''),
            'position_match': cv_data.get('position_match', {}),
            'generated_questions': cv_data.get('generated_questions', []),
            'strengths': cv_data.get('strengths', []),
            'areas_to_improve': cv_data.get('areas_to_improve', []),
            'suggested_questions': cv_data.get('suggested_questions', [])
        }
        
        # Quiz oluşturma seçeneği göster
        result['show_quiz_option'] = True
        
        # Log başarılı analizi
        logging.info(f"CV analizi başarılı: {cv_file.filename} - {position}")
        
        return jsonify(result)
        
    except Exception as e:
        logging.error(f"CV analizi hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': f"CV analiz edilirken bir hata oluştu: {str(e)}"
        })
    
def start_file_watcher():
    """Dosya izleme sistemini başlat - Sadece terminal log'ları için"""
    try:
        class InterviewFileHandler(FileSystemEventHandler):
            def on_created(self, event):
                if not event.is_directory:
                    if event.src_path.endswith('.json'):
                        print(f"\n[+] Yeni mülakat dosyası oluşturuldu: {os.path.basename(event.src_path)}")
                    elif event.src_path.endswith('.pdf'):
                        print(f"\n[+] Yeni rapor oluşturuldu: {os.path.basename(event.src_path)}")
                    
            def on_modified(self, event):
                if not event.is_directory and event.src_path.endswith('.json'):
                    print(f"\n[*] Mülakat dosyası güncellendi: {os.path.basename(event.src_path)}")
                    
            def on_deleted(self, event):
                if not event.is_directory:
                    if event.src_path.endswith('.json'):
                        print(f"\n[-] Mülakat dosyası silindi: {os.path.basename(event.src_path)}")
                    elif event.src_path.endswith('.pdf'):
                        print(f"\n[-] Rapor silindi: {os.path.basename(event.src_path)}")

        # İzlenecek dizinler
        paths_to_watch = ['interviews', 'reports']
        
        # Dizinleri oluştur
        for path in paths_to_watch:
            if not os.path.exists(path):
                os.makedirs(path)
                print(f"\n[+] {path} dizini oluşturuldu")
        
        # Observer oluştur ve başlat
        observer = Observer()
        for path in paths_to_watch:
            observer.schedule(InterviewFileHandler(), path, recursive=False)
        
        observer.start()
        print("\n[+] Dosya izleme sistemi başlatıldı")
        print("[*] interviews/ ve reports/ dizinleri izleniyor...")
        
        return observer
        
    except Exception as e:
        print(f"\n[!] Dosya izleme sistemi başlatılamadı: {str(e)}")
        return None

async def get_realtime_token():
    """OpenAI API anahtarını döndür"""
    try:
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            logger.error("OPENAI_API_KEY bulunamadı")
            return None
            
        return api_key
                
    except Exception as e:
        logger.error(f"Token alma hatası: {str(e)}")
        return None

@app.route('/get_interview_token', methods=['POST'])
async def get_interview_token():
    try:
        data = request.get_json()
        interview_code = data.get('code')
        
        if not interview_code:
            return jsonify({
                'success': False,
                'error': 'Mülakat kodu gerekli'
            })
            
        token = await get_realtime_token()
        if not token:
            return jsonify({
                'success': False,
                'error': 'Token alınamadı'
            })
            
        # Benzersiz bir session_id oluştur
        session_id = secrets.token_urlsafe(16)
            
        return jsonify({
            'success': True,
            'token': token,
            'session_id': session_id
        })
        
    except Exception as e:
        logger.error(f"Mülakat token hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Token alınırken bir hata oluştu'
        })

@app.route('/realtime_chat', methods=['POST'])
async def realtime_chat():
    try:
        data = request.get_json()
        message = data.get('message')
        session_id = data.get('session_id')
        interview_code = data.get('interview_code')
        
        if not message or not session_id or not interview_code:
            return jsonify({
                'success': False,
                'error': 'Mesaj, session_id ve interview_code gerekli'
            }), 400
            
        # JSON dosyasından mülakat verilerini al
        json_path = os.path.join('interviews', f'{interview_code}.json')
        if not os.path.exists(json_path):
            return jsonify({
                'success': False,
                'error': 'Mülakat bilgileri bulunamadı'
            }), 404
            
        with open(json_path, 'r', encoding='utf-8') as f:
            interview_data = json.load(f)
            
        # Konuşma geçmişini al veya oluştur
        conversation_history = interview_data.get('conversation_history', [])
        current_question_index = interview_data.get('current_question_index', 0)
        
        # Sistem promptunu hazırla
        system_prompt = f"""Sen deneyimli ve empatik bir İK uzmanısın. Şu anda {interview_data.get('candidate_name')} ile {interview_data.get('position')} pozisyonu için mülakat yapıyorsun.

        Mülakat soruları:
        {json.dumps(interview_data.get('questions', []), indent=2, ensure_ascii=False)}

        Şu anki soru indeksi: {current_question_index}
        
        Konuşma geçmişi:
        {json.dumps(conversation_history, indent=2, ensure_ascii=False)}

        Önemli Kurallar:
        1. Adayın hazır olma durumunu mutlaka kontrol et:
           - "Hazır değilim" veya benzeri bir yanıt gelirse:
             * "Anlıyorum, acele etmeyelim. Hazır olduğunuzda başlayalım. Kendinizi rahat hissettiğinizde bana haber verebilirsiniz."
             * "Biraz daha zaman ister misiniz? Ben buradayım, hazır olduğunuzda başlayabiliriz."
           - Aday hazır olduğunu belirtene kadar bir sonraki soruya geçme
        
        2. Soru tekrarı isteklerini dikkatle dinle:
           - "Anlamadım", "Tekrar eder misiniz" gibi ifadelerde:
             * "Tabii ki, soruyu farklı bir şekilde açıklayayım..."
             * "Elbette, şöyle sorayım..."
           - Soruyu farklı kelimelerle, daha açıklayıcı bir şekilde tekrar et
           - Aday anlayana kadar bir sonraki soruya geçme
        
        3. Her yanıttan sonra:
           - Adayın cevabını anladığını göster
           - Kısa ve samimi bir yorum yap
           - Eğer cevap yetersizse, nazikçe detay iste
           - Aday hazırsa ve cevap tamamsa, bir sonraki konuya doğal bir geçiş yap
        
        4. Doğal konuşma örnekleri:
           - "Bu konudaki deneyimlerinizi dinlemek çok değerli. Peki, [sonraki konu] hakkında ne düşünüyorsunuz?"
           - "Anlıyorum, yaklaşımınız ilginç. İsterseniz şimdi biraz da [yeni konu] hakkında konuşalım..."
           - "Bu tecrübeleriniz etkileyici. Başka bir konuya geçmeden önce eklemek istediğiniz bir şey var mı?"
        
        5. Eğer aday gergin veya stresli görünüyorsa:
           - "Kendinizi rahat hissetmeniz çok önemli. Acele etmeyelim..."
           - "Bu sadece bir sohbet, kendinizi baskı altında hissetmeyin..."
           - "İsterseniz biraz ara verebiliriz..."
        
        6. Mülakat bitişi:
           - Tüm sorular tamamlandığında nazik bir kapanış yap
           - "Paylaştığınız değerli bilgiler için teşekkür ederim. Görüşmemizi burada sonlandıralım. Size en kısa sürede dönüş yapacağız. MÜLAKAT_BİTTİ"
        
        En önemli nokta: Bu bir robot-insan konuşması değil, iki insan arasında geçen doğal bir sohbet olmalı. Adayın her tepkisine ve ihtiyacına uygun şekilde yanıt ver. Acele etme, adayın rahat hissetmesini sağla."""

        # OpenAI API'ye istek gönder
        try:
            completion = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: openai_client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        *conversation_history[-5:],  # Son 5 mesajı kullan
                        {"role": "user", "content": message}
                    ],
                    temperature=0.9,
                    max_tokens=250
                )
            )
            
            # Yanıtı al
            gpt_response = completion.choices[0].message.content
            is_interview_ended = "MÜLAKAT_BİTTİ" in gpt_response
            
            # Konuşma geçmişini güncelle
            conversation_history.append({"role": "user", "content": message})
            conversation_history.append({"role": "assistant", "content": gpt_response})
            
            # Soru indeksini güncelle (eğer bir soru sorulduysa)
            if "?" in gpt_response and not is_interview_ended:
                current_question_index += 1
            
            # JSON dosyasını güncelle
            interview_data['conversation_history'] = conversation_history
            interview_data['current_question_index'] = current_question_index
            
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(interview_data, f, ensure_ascii=False, indent=2)
            
            if is_interview_ended:
                try:
                    # Rapor oluşturma işlemini başlat
                    await generate_interview_report(interview_code, interview_data)
                    logger.info(f"Rapor oluşturma başarılı: {interview_code}")
                except Exception as e:
                    logger.error(f"Rapor oluşturma hatası: {str(e)}")
                # MÜLAKAT_BİTTİ kelimesini çıkar
                gpt_response = gpt_response.replace("MÜLAKAT_BİTTİ", "")
            
            return jsonify({
                'success': True,
                'text': gpt_response,
                'interview_ended': is_interview_ended
            })
            
        except Exception as e:
            logger.error(f"OpenAI API hatası: {str(e)}")
            return jsonify({
                'success': False,
                'error': f'OpenAI API hatası: {str(e)}'
            }), 500
            
    except Exception as e:
        logger.error(f"Gerçek zamanlı sohbet hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

async def generate_interview_report(interview_code, interview_data):
    try:
        # JSON dosyasını kontrol et
        json_path = os.path.join('interviews', f'{interview_code}.json')
        if not os.path.exists(json_path):
            raise ValueError(f"Mülakat dosyası bulunamadı: {interview_code}")
            
        with open(json_path, 'r', encoding='utf-8') as f:
            interview_json = json.load(f)
            
        # Webhook URL'sini kontrol et
        webhook_rapor_url = interview_json.get('webhook_rapor_url', WEBHOOK_RAPOR_URL)
        
        # PDF oluştur
        pdf_path = os.path.join('reports', f'{interview_code}.pdf')
        
        # Değerlendirme verilerini hazırla
        evaluation_data = {
            "teknik_yetkinlik": interview_data.get('teknik_yetkinlik', 0),
            "iletisim_becerileri": interview_data.get('iletisim_becerileri', 0),
            "problem_cozme": interview_data.get('problem_cozme', 0),
            "genel_degerlendirme": interview_data.get('genel_degerlendirme', ''),
            "guclu_yonler": interview_data.get('guclu_yonler', []),
            "gelisim_alanlari": interview_data.get('gelisim_alanlari', [])
        }
        
        # Webhook'a gönderilecek veriyi hazırla
        webhook_data = {
            "mulakat_kodu": interview_code,
            "aday_bilgileri": {
                "isim": interview_data.get('candidate_name'),
                "pozisyon": interview_data.get('position'),
                "tarih": datetime.now().isoformat()
            },
            "degerlendirme": evaluation_data,
            "rapor_dosyasi": pdf_path
        }
        
        # Webhook'a gönder
        webhook_response = requests.post(
            webhook_rapor_url,
            json=webhook_data,
            headers={'Content-Type': 'application/json'}
        )
        
        if webhook_response.status_code != 200:
            logger.error(f"Webhook hatası: {webhook_response.status_code} - {webhook_response.text}")
        
        # Mülakat durumunu güncelle
        interview_json['status'] = 'completed'
        interview_json['report_path'] = pdf_path
        interview_json['evaluation_data'] = evaluation_data
        
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(interview_json, f, ensure_ascii=False, indent=2)
        
        return jsonify({
            'success': True,
            'message': 'Mülakat raporu oluşturuldu ve webhook\'a gönderildi',
            'report_path': pdf_path
        })
        
    except Exception as e:
        logger.error(f"Rapor oluşturma hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/end_interview', methods=['POST'])
async def end_interview():
    try:
        data = request.get_json()
        interview_code = data.get('interview_code')
        conversation_history = data.get('conversation_history', [])
        
        if not interview_code:
            return jsonify({
                'success': False,
                'error': 'Mülakat kodu gerekli'
            }), 400
            
        # Rapor oluşturma işlemini başlat
        await generate_interview_report(interview_code, conversation_history)
        
        return jsonify({
            'success': True,
            'message': 'Mülakat sonlandırıldı ve rapor oluşturma işlemi başlatıldı'
        })
        
    except Exception as e:
        logger.error(f"Mülakat sonlandırma hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# window.onbeforeunload event'i için endpoint
@app.route('/save_interview_state', methods=['POST'])
def save_interview_state():
    try:
        data = request.json
        code = data.get('code')
        conversation_history = data.get('conversation_history', [])
        ended = data.get('ended', False)
        
        if not code:
            return jsonify({'success': False, 'error': 'Mülakat kodu gereklidir'}), 400
            
        # JSON dosyasını kontrol et
        json_path = os.path.join('interviews', f'{code}.json')
        if not os.path.exists(json_path):
            return jsonify({'success': False, 'error': 'Mülakat bulunamadı'}), 404
            
        # JSON dosyasını oku
        with open(json_path, 'r', encoding='utf-8') as f:
            interview_data = json.load(f)
            
        # Verileri güncelle
        interview_data['conversation_history'] = conversation_history
        if ended:
            interview_data['ended'] = True
            interview_data['status'] = 'completed'
            interview_data['ended_at'] = datetime.now().isoformat()
            
        # Verileri kaydet
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(interview_data, f, ensure_ascii=False, indent=2)
            
        return jsonify({'success': True})
        
    except Exception as e:
        logger.error(f"Mülakat durumu kaydetme hatası: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/get_speech', methods=['GET'])
async def get_speech():
    try:
        text = request.args.get('text')
        if not text:
            return 'Text parameter is required', 400

        # Geçici ses dosyası için benzersiz bir isim oluştur
        temp_filename = f"temp/speech_{int(time.time())}_{random.randint(100000, 999999)}.mp3"
        
        try:
            # OpenAI TTS ile sesi oluştur - context manager kullanmadan
            tts_response = openai_client.audio.speech.create(
                model="tts-1",
                voice="nova",
                input=text
            )
            
            # Ses dosyasını kaydet
            tts_response.stream_to_file(temp_filename)
            
            # Dosyayı oku ve yanıt olarak gönder
            def generate():
                with open(temp_filename, 'rb') as f:
                    while chunk := f.read(8192):
                        yield chunk
                # Dosyayı sil
                if os.path.exists(temp_filename):
                    os.remove(temp_filename)
            
            return Response(
                generate(),
                mimetype='audio/mpeg',
                headers={
                    'Content-Disposition': 'inline',
                    'Cache-Control': 'no-cache'
                }
            )
            
        except Exception as e:
            logger.error(f"TTS hatası: {str(e)}")
            if os.path.exists(temp_filename):
                os.remove(temp_filename)
            return jsonify({
                'success': False,
                'error': 'Ses oluşturulamadı'
            }), 500
            
    except Exception as e:
        logger.error(f"Ses endpoint hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# Pozisyona göre çoktan seçmeli soru oluşturma fonksiyonu
@app.route('/create_multiple_choice', methods=['POST'])
def create_multiple_choice():
    """Çoktan seçmeli sorular oluştur"""
    try:
        position = request.json.get('position', '')
        question_count = int(request.json.get('question_count', 10))
        quiz_type = request.json.get('quiz_type', 'position')
        cv_data = request.json.get('cv_data', {})
        
        if not position:
            return jsonify({
                'success': False,
                'error': 'Pozisyon bilgisi gereklidir'
            })
        
        # Mülakat kodu oluştur (eğer sağlanmamışsa)
        interview_code = request.json.get('interview_code')
        if not interview_code:
            # Yeni ve belirgin bir mülakat kodu oluştur
            interview_code = generate_interview_code()
            
            # Yeni mülakat verisi oluştur
            interview_data = {
                'code': interview_code,
                'candidate_name': request.json.get('candidate_name', 'İsimsiz Aday'),
                'position': position,
                'status': 'pending',
                'created_at': datetime.now().isoformat(),
                'cv_data': cv_data,
                'quiz_created': True,  # Quiz oluşturulduğunu belirt
                'interview_started': False  # Mülakat henüz başlamadı
            }
            
            # Mülakat verisini kaydet
            save_interview_data(interview_data, interview_code)
            logger.info(f"Quiz için yeni mülakat oluşturuldu: {interview_code}")
        else:
            # Mevcut mülakatı güncelle - quiz oluşturulduğunu belirt
            try:
                json_path = os.path.join('interviews', f'{interview_code}.json')
                if os.path.exists(json_path):
                    with open(json_path, 'r', encoding='utf-8') as f:
                        interview_data = json.load(f)
                    
                    # Quiz oluşturuldu olarak işaretle
                    interview_data['quiz_created'] = True
                    
                    # Güncellenen veriyi kaydet
                    with open(json_path, 'w', encoding='utf-8') as f:
                        json.dump(interview_data, f, ensure_ascii=False, indent=2)
            except Exception as update_error:
                logger.error(f"Mülakat verisi güncelleme hatası: {str(update_error)}")
                # Bu hata quiz oluşturmayı engellememelidir

        # Giriş mesajını hazırla
        system_message = """
        Sen bir teknik mülakat uzmanısın. Verilen pozisyon ve CV bilgileri doğrultusunda, 
        adayın teknik yetkinliklerini ölçmek için çoktan seçmeli sorular hazırlayacaksın.
        
        Her soru aşağıdaki formatta olmalıdır:
        
        1. [Soru metni]
        
        A. [Seçenek metni]
        B. [Seçenek metni]
        C. [Seçenek metni]
        D. [Seçenek metni]
        
        Doğru cevap: [A, B, C veya D]
        
        Açıklama: [Cevabın açıklaması]
        """
        
        # CV bilgilerini içeren bağlam oluştur
        cv_context = ""
        if cv_data and (quiz_type == 'cv' or quiz_type == 'both'):
            skills = cv_data.get('skills', [])
            experiences = cv_data.get('experience', [])
            education = cv_data.get('education', [])
            
            if skills:
                cv_context += "Adayın becerileri: " + ", ".join(skills) + "\n\n"
            
            if experiences:
                cv_context += "Adayın deneyimleri:\n"
                for exp in experiences:
                    cv_context += f"- {exp.get('title', 'Pozisyon')} at {exp.get('company', 'Şirket')}: {exp.get('description', '')}\n"
                cv_context += "\n"
            
            if education:
                cv_context += "Adayın eğitimi:\n"
                for edu in education:
                    cv_context += f"- {edu.get('degree', 'Derece')} in {edu.get('field', 'Alan')}, {edu.get('institution', 'Kurum')}\n"
                cv_context += "\n"
        
        # Kullanıcı komutunu hazırla
        if quiz_type == 'position':
            user_message = f"""
            {position} pozisyonu için {question_count} adet çoktan seçmeli teknik soru oluştur.
            Her soru için 4 şık (A, B, C, D) hazırla ve doğru cevabı belirt.
            Her soru için bir açıklama ve doğru cevabın neden doğru olduğuna dair kısa bir bilgi ekle.
            
            Lütfen aşağıdaki formatı takip et:
            1. [Soru metni]
            
            A. [Seçenek metni]
            B. [Seçenek metni]
            C. [Seçenek metni]
            D. [Seçenek metni]
            
            Doğru cevap: [A, B, C veya D]
            
            Açıklama: [Cevabın açıklaması]
            """
        elif quiz_type == 'cv':
            user_message = f"""
            Aşağıdaki CV bilgilerine dayanarak, adayın beceri ve deneyimlerini test etmek için {question_count} adet çoktan seçmeli teknik soru oluştur.
            
            {cv_context}
            
            Her soru için 4 şık (A, B, C, D) hazırla ve doğru cevabı belirt.
            Her soru için bir açıklama ve doğru cevabın neden doğru olduğuna dair kısa bir bilgi ekle.
            
            Lütfen aşağıdaki formatı takip et:
            1. [Soru metni]
            
            A. [Seçenek metni]
            B. [Seçenek metni]
            C. [Seçenek metni]
            D. [Seçenek metni]
            
            Doğru cevap: [A, B, C veya D]
            
            Açıklama: [Cevabın açıklaması]
            """
        elif quiz_type == 'both':
            user_message = f"""
            {position} pozisyonu ve aşağıdaki CV bilgilerine dayanarak, adayın teknik yetkinliklerini ölçmek için {question_count} adet çoktan seçmeli soru oluştur.
            
            {cv_context}
            
            Sorular hem pozisyon gereksinimleri hem de adayın becerileri ile ilgili olmalıdır.
            Her soru için 4 şık (A, B, C, D) hazırla ve doğru cevabı belirt.
            Her soru için bir açıklama ve doğru cevabın neden doğru olduğuna dair kısa bir bilgi ekle.
            
            Lütfen aşağıdaki formatı takip et:
            1. [Soru metni]
            
            A. [Seçenek metni]
            B. [Seçenek metni]
            C. [Seçenek metni]
            D. [Seçenek metni]
            
            Doğru cevap: [A, B, C veya D]
            
            Açıklama: [Cevabın açıklaması]
            """
        
        # OpenAI API çağrısı
        try:
            logger.info(f"OpenAI API'ye istek gönderiliyor - pozisyon: {position}, soru sayısı: {question_count}")
            completion = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": user_message}
                ],
                temperature=0.5
            )
            
            # Yanıtı analiz et
            content = completion.choices[0].message.content
            logger.info(f"OpenAI yanıtı alındı - uzunluk: {len(content)}")
        except Exception as api_error:
            logger.error(f"OpenAI API hatası: {str(api_error)}")
            return jsonify({
                'success': False,
                'error': f"OpenAI API hatası: {str(api_error)}"
            })
        
        # Soruları ayrıştır
        questions = []
        current_question = {}
        question_id = 0
        lines = content.split('\n')
        line_count = len(lines)
        
        logger.info(f"Toplam {line_count} satır ayrıştırılacak")
        
        try:
            i = 0
            while i < line_count:
                line = lines[i].strip()
                
                # Yeni soru başlangıcı
                if re.match(r'^(\d+)\.\s+', line):
                    if current_question and 'question' in current_question:
                        questions.append(current_question)
                    
                    question_id += 1
                    question_match = re.match(r'^(\d+)\.\s+(.*)', line)
                    if question_match:
                        question_text = question_match.group(2)
                        current_question = {
                            'id': str(question_id),
                            'question': question_text,
                            'options': {},
                            'correct_answer': '',
                            'explanation': ''
                        }
                
                # Şık ayrıştırma
                elif re.match(r'^[A-D]\.\s+', line):
                    option_match = re.match(r'^([A-D])\.\s+(.*)', line)
                    if option_match and current_question:
                        option_letter = option_match.group(1)
                        option_text = option_match.group(2)
                        if 'options' in current_question:
                            current_question['options'][option_letter] = option_text
                        else:
                            logger.warning(f"Soru ID {question_id} için options anahtarı yok")
                
                # Doğru cevap
                elif "doğru cevap" in line.lower() or "cevap:" in line.lower():
                    answer_match = re.search(r'[A-D]', line)
                    if answer_match and current_question:
                        current_question['correct_answer'] = answer_match.group(0)
                
                # Açıklama
                elif "açıklama" in line.lower() or "explanation" in line.lower():
                    if current_question:
                        explanation_parts = []
                        if ":" in line:
                            explanation_parts.append(line.split(":", 1)[1].strip())
                        else:
                            explanation_parts.append(line)
                        
                        # Sonraki satırları açıklamaya ekle
                        j = i + 1
                        while j < line_count:
                            next_line = lines[j].strip()
                            if not next_line or re.match(r'^(\d+)\.\s+', next_line) or re.match(r'^[A-D]\.\s+', next_line):
                                break
                            explanation_parts.append(next_line)
                            j += 1
                        
                        current_question['explanation'] = " ".join(explanation_parts).strip()
                
                i += 1
        
            # Son soruyu da ekle
            if current_question and 'question' in current_question:
                questions.append(current_question)
            
            logger.info(f"GPT tarafından {len(questions)} soru oluşturuldu")
            
            # Hiç soru yoksa hata döndür
            if len(questions) == 0:
                logger.error(f"GPT yanıtından hiç soru oluşturulamadı. API yanıtı: {content}")
                return jsonify({
                    'success': False,
                    'error': 'Sorular oluşturulamadı. Lütfen tekrar deneyin veya farklı bir pozisyon girin.'
                })
                
            # Tüm soruların gerekli alanlarını kontrol et
            valid_questions = []
            for q in questions:
                if ('question' in q and 'options' in q and len(q['options']) > 0 and 
                    'correct_answer' in q and q['correct_answer']):
                    valid_questions.append(q)
                else:
                    logger.warning(f"Geçersiz soru formatı: {q}")
            
            # Soruları dosyaya kaydet
            quiz_filename = f"quiz_{datetime.now().strftime('%Y%m%d%H%M%S')}_{interview_code}.json"
            quiz_path = os.path.join(QUESTIONS_DIR, quiz_filename)
            
            quiz_data = {
                'position': position,
                'interview_code': interview_code,
                'quiz_type': quiz_type,
                'cv_context': cv_context if quiz_type in ['cv', 'both'] else "",
                'created_at': datetime.now().isoformat(),
                'questions': valid_questions
            }
            
            with open(quiz_path, 'w', encoding='utf-8') as f:
                json.dump(quiz_data, f, ensure_ascii=False, indent=2)
            
            return jsonify({
                'success': True,
                'message': f"{len(valid_questions)} adet çoktan seçmeli soru oluşturuldu.",
                'quiz_url': f"/quiz/{interview_code}",
                'interview_code': interview_code
            })
            
        except Exception as parse_error:
            logger.error(f"Soru ayrıştırma hatası: {str(parse_error)}")
            logger.error(f"İçerik: {content[:200]}...")
            return jsonify({
                'success': False,
                'error': f"Sorular ayrıştırılırken bir hata oluştu: {str(parse_error)}"
            })
        
    except Exception as e:
        logging.error(f"Çoktan seçmeli soru oluşturma hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': f"Soru oluşturulurken bir hata oluştu: {str(e)}"
        })

# Çoktan seçmeli sınav sayfasını göster
@app.route('/quiz/<interview_code>')
def show_quiz(interview_code):
    """Çoktan seçmeli sınavı göster"""
    try:
        # Eğer gelen kod garip bir format (JSON dökümü gibi) içeriyorsa temizle
        if isinstance(interview_code, str) and ('{' in interview_code or '}' in interview_code):
            logger.warning(f"Geçersiz quiz kodu temizlendi: {interview_code}")
            return render_template('error.html', message="Geçersiz sınav kodu. Lütfen doğru bir kod kullanınız.")
        
        # İlgili soru dosyasını bul
        quiz_file = None
        for file in os.listdir(QUESTIONS_DIR):
            if file.endswith(f"_{interview_code}.json"):
                quiz_file = file
                break
        
        if not quiz_file:
            return render_template('error.html', message="Sınav bulunamadı. Lütfen kod doğruluğunu kontrol ediniz.")
        
        # Soru dosyasını oku
        with open(os.path.join(QUESTIONS_DIR, quiz_file), 'r', encoding='utf-8') as f:
            quiz_data = json.load(f)
        
        # Şablonu göster
        return render_template('quiz.html', 
                              interview_code=interview_code,
                              position=quiz_data.get('position', ''),
                              questions=quiz_data.get('questions', []))
    
    except Exception as e:
        logging.error(f"Çoktan seçmeli sınav gösterme hatası: {str(e)}")
        return render_template('error.html', message=f"Sınav gösterilirken bir hata oluştu: {str(e)}")

@app.route('/save_photo', methods=['POST'])
def save_photo():
    try:
        data = request.json
        code = data.get('code')
        photo_data = data.get('photo')
        timestamp = data.get('timestamp')
        
        if not code or not photo_data:
            return jsonify({'success': False, 'error': 'Mülakat kodu ve fotoğraf verisi gereklidir'}), 400
            
        # Geçersiz mülakat kodunu temizle
        if isinstance(code, str) and ('{' in code or '}' in code):
            logger.warning(f"Geçersiz fotoğraf kaydetme mülakat kodu: {code}")
            return jsonify({'success': False, 'error': 'Geçersiz mülakat kodu formatı'}), 400
            
        # Quiz dosyası kontrolü
        quiz_found = False
        for file in os.listdir(QUESTIONS_DIR):
            if file.endswith(f"_{code}.json"):
                quiz_found = True
                break
                
        # Eğer quiz dosyası varsa, fotoğraflar için quiz kodu kullanılabilir
        if quiz_found:
            # Quiz fotoğraflarını saklamak için klasör oluştur
            photos_dir = os.path.join('interviews', f'quiz_{code}_photos')
            os.makedirs(photos_dir, exist_ok=True)
            
            # Base64 formatındaki fotoğrafı işle
            try:
                # Başlangıç kısmını (data:image/jpeg;base64,) kaldır
                if ',' in photo_data:
                    image_data = photo_data.split(',')[1]
                else:
                    image_data = photo_data
                
                # Fotoğrafa benzersiz bir ad ver
                photo_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}.jpg"
                photo_path = os.path.join(photos_dir, photo_filename)
                
                # Fotoğrafı kaydet
                with open(photo_path, 'wb') as f:
                    f.write(base64.b64decode(image_data))
                
                return jsonify({'success': True})
            except Exception as e:
                logger.error(f"Quiz fotoğraf işleme hatası: {str(e)}")
                return jsonify({'success': False, 'error': f'Fotoğraf işlenirken hata: {str(e)}'}), 500
            
        # JSON dosyasını kontrol et
        json_path = os.path.join('interviews', f'{code}.json')
        if not os.path.exists(json_path):
            return jsonify({'success': False, 'error': 'Mülakat bulunamadı'}), 404
        
        # Base64 formatındaki fotoğrafı işle
        # Başlangıç kısmını (data:image/jpeg;base64,) kaldır
        if ',' in photo_data:
            image_data = photo_data.split(',')[1]
        else:
            image_data = photo_data
        
        # Fotoğrafları saklamak için klasör oluştur
        photos_dir = os.path.join('interviews', f'{code}_photos')
        os.makedirs(photos_dir, exist_ok=True)
        
        # Fotoğrafa benzersiz bir ad ver
        photo_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}.jpg"
        photo_path = os.path.join(photos_dir, photo_filename)
        
        # Fotoğrafı kaydet
        with open(photo_path, 'wb') as f:
            f.write(base64.b64decode(image_data))
            
        # JSON dosyasını güncelle (fotoğraf bilgilerini ekle)
        with open(json_path, 'r', encoding='utf-8') as f:
            interview_data = json.load(f)
            
        # Fotoğraf bilgilerini ekle
        if 'photos' not in interview_data:
            interview_data['photos'] = []
            
        interview_data['photos'].append({
            'filename': photo_filename,
            'timestamp': timestamp,
            'path': os.path.join(f'{code}_photos', photo_filename)
        })
        
        # Güncellenen veriyi kaydet
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(interview_data, f, ensure_ascii=False, indent=2)
            
        return jsonify({'success': True})
        
    except Exception as e:
        logger.error(f"Fotoğraf kaydetme hatası: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/submit_quiz/<interview_code>', methods=['POST'])
def submit_quiz(interview_code):
    """Sınav cevaplarını kaydet ve sonuçları göster"""
    try:
        # Gelen kodda JSON formatı veya geçersiz karakterler varsa temizle
        if isinstance(interview_code, str) and ('{' in interview_code or '}' in interview_code):
            logger.warning(f"Geçersiz submit_quiz kodu temizlendi: {interview_code}")
            return jsonify({
                'success': False,
                'error': "Geçersiz sınav kodu. Lütfen doğru bir kod kullanınız."
            })
            
        # Cevapları al
        answers = request.get_json()
        
        # İlgili soru dosyasını bul
        quiz_file = None
        for file in os.listdir(QUESTIONS_DIR):
            if file.endswith(f"_{interview_code}.json"):
                quiz_file = file
                break
        
        if not quiz_file:
            return jsonify({
                'success': False,
                'error': "Sınav bulunamadı. Lütfen kod doğruluğunu kontrol ediniz."
            })
        
        # Soru dosyasını oku
        with open(os.path.join(QUESTIONS_DIR, quiz_file), 'r', encoding='utf-8') as f:
            quiz_data = json.load(f)
        
        # Sonuçları hesapla
        correct_count = 0
        results = []
        
        # Doğru cevapları kontrol et
        all_answers = [q.get('correct_answer', '').strip().upper() for q in quiz_data.get('questions', [])]
        unique_answers = set(all_answers)
        
        # Eğer tüm doğru cevaplar "D" ise, quiz dosyasının kendisini düzelt
        if len(unique_answers) == 1 and 'D' in unique_answers and len(all_answers) > 3:
            logger.warning(f"Bozuk quiz dosyası tespit edildi - tüm cevaplar 'D'. Doğru cevapları rastgele düzeltiyorum.")
            
            # Seçenekler listesi
            option_letters = ['A', 'B', 'C', 'D']
            
            # Her soruya rastgele doğru cevap ata
            for question in quiz_data.get('questions', []):
                # Mevcut doğru cevap
                current_answer = question.get('correct_answer', 'D').strip().upper()
                if current_answer == 'D':
                    # Rastgele yeni bir cevap seç (tüm cevaplar D olmayacak şekilde)
                    import random
                    new_answer = random.choice(option_letters)
                    question['correct_answer'] = new_answer
                    logger.info(f"Soru {question.get('id')} için doğru cevap 'D'den '{new_answer}'a değiştirildi")
            
            # Düzeltilmiş quiz dosyasını kaydet
            quiz_path = os.path.join(QUESTIONS_DIR, quiz_file)
            with open(quiz_path, 'w', encoding='utf-8') as f:
                json.dump(quiz_data, f, ensure_ascii=False, indent=2)
                
            logger.info(f"Düzeltilmiş quiz dosyası kaydedildi: {quiz_path}")
        
        # Normal değerlendirme
        for question in quiz_data.get('questions', []):
            question_id = question.get('id')
            user_answer = answers.get(str(question_id))
            correct_answer = question.get('correct_answer')
            
            # Veri tiplerini düzenle - Her ikisini de string'e çevir ve büyük harfe dönüştür
            if user_answer is not None:
                user_answer = str(user_answer).strip().upper()
            else:
                user_answer = ""
                
            if correct_answer is not None:
                correct_answer = str(correct_answer).strip().upper()
            else:
                correct_answer = ""
            
            # Karşılaştırma
            is_correct = user_answer == correct_answer
            
            # Sonuç kontrolü
            logger.info(f"Karşılaştırma sonucu: {is_correct} (user_answer='{user_answer}', correct_answer='{correct_answer}')")
            
            if is_correct:
                correct_count += 1
                
            # Sonuçları kaydet - JSON dosyası için tam veri
            detailed_result = {
                'id': question_id,
                'question': question.get('question'),
                'user_answer': user_answer,
                'correct_answer': correct_answer,
                'is_correct': is_correct,
                'explanation': question.get('explanation')
            }
            
            # Kullanıcıya gösterilecek sonuçlar - doğru cevapları içermez
            user_result = {
                'id': question_id,
                'question': question.get('question'),
                'user_answer': user_answer,
                'is_correct': is_correct,
                'explanation': "" if is_correct else "Bu soruya yanlış cevap verdiniz."
            }
            
            # İki farklı sonuç listesi tut
            results.append(user_result)  # Kullanıcıya gösterilecek
        
        # Sonuçları kaydet
        total_questions = len(quiz_data.get('questions', []))
        if total_questions > 0:
            score_percentage = int((correct_count / total_questions) * 100)
        else:
            score_percentage = 0
        
        # Timestamp ekle
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Başarısız olan sorular ve başarılı olan sorular
        failed_questions = [q for q in quiz_data.get('questions', []) if answers.get(str(q.get('id'))) != q.get('correct_answer')]
        passed_questions = [q for q in quiz_data.get('questions', []) if answers.get(str(q.get('id'))) == q.get('correct_answer')]
        
        # Yanlış yapılan soru konuları
        failed_topics = [q.get('question', '').split()[0:3] for q in failed_questions]
        failed_topics_str = ", ".join([" ".join(topic) for topic in failed_topics])
        
        # Doğru cevapların yüzdesi ve analizi
        analysis = {
            'correct_percentage': score_percentage,
            'correct_count': correct_count,
            'total_questions': total_questions,
            'failed_topics': failed_topics_str if failed_topics else "Tüm konularda başarılı",
            'performance': "Çok iyi" if score_percentage >= 90 else 
                          "İyi" if score_percentage >= 75 else 
                          "Orta" if score_percentage >= 50 else 
                          "Geliştirilebilir"
        }
        
        # Tam sonuç verilerini kaydet (doğru cevapları içerir, sadece JSON'da saklanır)
        detailed_result_data = {
            'interview_code': interview_code,
            'position': quiz_data.get('position', ''),
            'total_questions': total_questions,
            'correct_count': correct_count,
            'percentage': score_percentage,
            'timestamp': timestamp,
            'analysis': analysis,
            'detailed_results': [detailed_result for question in quiz_data.get('questions', []) 
                                for detailed_result in [
                                    {
                                        'id': question.get('id'),
                                        'question': question.get('question'),
                                        'user_answer': answers.get(str(question.get('id')), ""),
                                        'correct_answer': question.get('correct_answer', ""),
                                        'is_correct': answers.get(str(question.get('id'))) == question.get('correct_answer'),
                                        'explanation': question.get('explanation', "")
                                    }
                                ]]
        }
        
        # Sonuç dosyasını kaydet
        result_filename = f"result_{interview_code}_{int(time.time())}.json"
        result_path = os.path.join(QUESTIONS_DIR, result_filename)
        with open(result_path, 'w', encoding='utf-8') as f:
            json.dump(detailed_result_data, f, ensure_ascii=False, indent=2)
        
        # Mülakat verisine quiz sonuçlarını ekle
        try:
            json_path = os.path.join('interviews', f'{interview_code}.json')
            if os.path.exists(json_path):
                with open(json_path, 'r', encoding='utf-8') as f:
                    interview_data = json.load(f)
                
                # Quiz sonuçlarını ekle
                interview_data['quiz_results'] = {
                    'timestamp': timestamp,
                    'score': {
                        'correct': correct_count,
                        'total': total_questions,
                        'percentage': score_percentage
                    },
                    'analysis': analysis,  # Analiz verilerini de ekle
                    'result_file': result_filename
                }
                
                # Rapor oluştur - quiz sonuçlarıyla
                if 'report_generated' not in interview_data or not interview_data['report_generated']:
                    try:
                        # Asenkron rapor oluşturmayı başlat
                        asyncio.run(generate_quiz_report(interview_code, interview_data))
                        interview_data['report_generated'] = True
                    except Exception as report_error:
                        logger.error(f"Quiz raporu oluşturma hatası: {str(report_error)}")
                
                # Güncellenen veriyi kaydet
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(interview_data, f, ensure_ascii=False, indent=2)
        except Exception as update_error:
            logger.error(f"Quiz sonuçlarını mülakat verisine ekleme hatası: {str(update_error)}")
            # Bu hata quiz sonuçlarının dönüşünü engellememelidir
        
        # Ana return ifadesi - tüm işlemler tamamlandığında sonuçları döndür
        # Kullanıcıya doğru cevapları göstermiyoruz, sadece kendi cevaplarının doğru/yanlış olduğunu söylüyoruz
        return jsonify({
            'success': True,
            'correct_count': correct_count,
            'total_questions': total_questions,
            'percentage': score_percentage,
            'results': results,  # Sadece kullanıcı sonuçlarını gönder (doğru cevapları içermez)
            'analysis': analysis['performance']  # Performans değerlendirmesini ekle
        })
            
    except Exception as e:
        logging.error(f"Çoktan seçmeli sınav sonuçlarını kaydetme hatası: {str(e)}")
        return jsonify({
            'success': False,
            'error': f"Sınav sonuçları kaydedilirken bir hata oluştu: {str(e)}"
        })

# Quiz raporu oluşturma fonksiyonu
async def generate_quiz_report(interview_code, interview_data):
    """Sadece quiz sonuçlarına dayalı rapor oluştur"""
    try:
        if not interview_code:
            logger.error("Rapor oluşturmak için mülakat kodu gereklidir")
            return False
            
        logger.info(f"Quiz raporu oluşturuluyor: {interview_code}")
        
        # Quiz sonuçlarını al
        quiz_results = interview_data.get('quiz_results', {})
        if not quiz_results:
            logger.error(f"Quiz raporu oluşturulamadı: Quiz sonuçları bulunamadı. Kod: {interview_code}")
            return False
        
        # Quiz analiz verilerini al
        quiz_analysis = quiz_results.get('analysis', {})
        
        # Sonuç dosyasından detaylı verileri al
        result_file = quiz_results.get('result_file')
        detailed_results = []
        if result_file:
            result_path = os.path.join(QUESTIONS_DIR, result_file)
            if os.path.exists(result_path):
                try:
                    with open(result_path, 'r', encoding='utf-8') as f:
                        result_data = json.load(f)
                        detailed_results = result_data.get('detailed_results', [])
                except Exception as e:
                    logger.error(f"Sonuç dosyası okuma hatası: {str(e)}")
        
        # Mülakat sonuçlarını kontrol et
        interview_results = interview_data.get('evaluation_results', {})
        has_interview_results = bool(interview_results)
            
        # Rapor verileri
        report_data = {
            'interview_code': interview_code,
            'candidate_name': interview_data.get('candidate_name', 'İsimsiz Aday'),
            'position': interview_data.get('position', 'Belirtilmemiş Pozisyon'),
            'date': datetime.now().strftime('%d.%m.%Y'),
            'quiz_score': quiz_results.get('score', {}).get('percentage', 0),
            'quiz_correct': quiz_results.get('score', {}).get('correct', 0),
            'quiz_total': quiz_results.get('score', {}).get('total', 0),
            'quiz_only': not has_interview_results,  # Sadece quiz ise True, mülakat da varsa False
            'quiz_analysis': quiz_analysis,  # Quiz analiz verilerini ekle
            'detailed_results': detailed_results  # Detaylı soru sonuçlarını ekle
        }
        
        # Mülakat sonuçları varsa, onları da ekle
        if has_interview_results:
            report_data.update({
                'technical_score': interview_results.get('technical_score', 0),
                'communication_score': interview_results.get('communication_score', 0),
                'confidence_score': interview_results.get('confidence_score', 0),
                'overall_score': interview_results.get('overall_score', 0)
            })
        else:
            # Mülakat sonuçları yoksa, quiz skorunu teknik skor olarak kullan
            report_data.update({
                'technical_score': quiz_results.get('score', {}).get('percentage', 0),
                'overall_score': quiz_results.get('score', {}).get('percentage', 0)
            })
        
        # PDF oluştur
        report_renderer = QuizReportRenderer(report_data)
        pdf_path = report_renderer.generate_pdf()
        
        # PDF dosya adını mülakat verisine ekle
        if pdf_path:
            interview_data['report_pdf'] = os.path.basename(pdf_path)
            interview_data['report_generated_at'] = datetime.now().isoformat()
            interview_data['report_status'] = 'completed'
            
            # Güncellenen veriyi kaydet
            json_path = os.path.join('interviews', f'{interview_code}.json')
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(interview_data, f, ensure_ascii=False, indent=2)
                
            logger.info(f"Quiz raporu başarıyla oluşturuldu: {pdf_path}")
            return True
        else:
            logger.error(f"Quiz raporu oluşturulamadı: PDF dosyası oluşturulamadı. Kod: {interview_code}")
            return False
            
    except Exception as e:
        logger.error(f"Quiz raporu oluşturma hatası: {str(e)}")
        return False

# Çoktan seçmeli sınav giriş sayfası
@app.route('/quiz_entry')
def quiz_entry():
    """Çoktan seçmeli sınav giriş sayfasını göster"""
    return render_template('quiz_entry.html')

# ReportRenderer sınıfı - PDF raporları oluşturmak için
class QuizReportRenderer:
    """Quiz PDF raporlarını oluşturmak için özel sınıf"""
    def __init__(self, report_data):
        self.report_data = report_data
        self.styles = getSampleStyleSheet()
        
    def generate_pdf(self):
        try:
            # PDF dosya yolu
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            interview_code = self.report_data.get('interview_code', 'unknown')
            pdf_filename = f"quiz_raporu_{interview_code}_{timestamp}.pdf"
            pdf_path = os.path.join('reports', pdf_filename)
            
            # PDF oluştur
            doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            story = []
            
            # Başlık
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=self.styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1  # Ortalanmış
            )
            story.append(Paragraph("Quiz Değerlendirme Raporu", title_style))
            story.append(Spacer(1, 12))
            
            # Aday Bilgileri
            info_style = ParagraphStyle(
                'Info',
                parent=self.styles['Normal'],
                fontSize=12,
                spaceAfter=6
            )
            story.append(Paragraph(f"Aday: {self.report_data.get('candidate_name', 'Belirtilmemiş')}", info_style))
            story.append(Paragraph(f"Pozisyon: {self.report_data.get('position', 'Belirtilmemiş')}", info_style))
            story.append(Paragraph(f"Tarih: {self.report_data.get('date', datetime.now().strftime('%d.%m.%Y'))}", info_style))
            story.append(Paragraph(f"Mülakat Kodu: {self.report_data.get('interview_code', 'Belirtilmemiş')}", info_style))
            story.append(Spacer(1, 20))
            
            # Quiz Sonuçları
            story.append(Paragraph("Quiz Sonuçları", self.styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # Quiz skoru ve istatistikleri
            quiz_score = self.report_data.get('quiz_score', 0)
            quiz_correct = self.report_data.get('quiz_correct', 0)
            quiz_total = self.report_data.get('quiz_total', 0)
            
            metrics_data = [
                ["Metrik", "Değer"],
                ["Doğru Cevaplar", f"{quiz_correct}"],
                ["Toplam Soru", f"{quiz_total}"],
                ["Quiz Skoru", f"{quiz_score}%"],
                ["Performans", self.report_data.get('quiz_analysis', {}).get('performance', 'Değerlendirilmedi')]
            ]
            
            t = Table(metrics_data, colWidths=[300, 100])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(t)
            story.append(Spacer(1, 20))
            
            # Quiz Analizi
            story.append(Paragraph("Quiz Analizi", self.styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # Quiz analizini ekle
            quiz_analysis = self.report_data.get('quiz_analysis', {})
            
            # Başarısız olunan konular
            failed_topics = quiz_analysis.get('failed_topics', "Veri yok")
            
            analysis_text = f"""Aday {quiz_total} sorudan {quiz_correct} tanesine doğru cevap vermiştir (%{quiz_score} başarı).

Performans değerlendirmesi: {quiz_analysis.get('performance', 'Veri yok')}

Eksik görülen konular: {failed_topics}"""
            
            story.append(Paragraph(analysis_text, self.styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Doğru/Yanlış Soruları Listele
            story.append(Paragraph("Soru Detayları", self.styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # Detaylı sonuçları göster
            detailed_results = self.report_data.get('detailed_results', [])
            if detailed_results:
                # Doğru cevaplar
                correct_questions = [r for r in detailed_results if r.get('is_correct', False)]
                if correct_questions:
                    story.append(Paragraph("Doğru Cevaplanan Sorular:", self.styles['Heading3']))
                    story.append(Spacer(1, 8))
                    
                    for i, result in enumerate(correct_questions, 1):
                        question = result.get('question', '')
                        correct_answer = result.get('correct_answer', '')
                        story.append(Paragraph(f"{i}. {question} (Doğru cevap: {correct_answer})", self.styles['Normal']))
                    
                    story.append(Spacer(1, 12))
                
                # Yanlış cevaplar
                wrong_questions = [r for r in detailed_results if not r.get('is_correct', False)]
                if wrong_questions:
                    story.append(Paragraph("Yanlış Cevaplanan Sorular:", self.styles['Heading3']))
                    story.append(Spacer(1, 8))
                    
                    for i, result in enumerate(wrong_questions, 1):
                        question = result.get('question', '')
                        user_answer = result.get('user_answer', '')
                        correct_answer = result.get('correct_answer', '')
                        explanation = result.get('explanation', '')
                        
                        story.append(Paragraph(f"{i}. {question}", self.styles['Normal']))
                        story.append(Paragraph(f"   Adayın cevabı: {user_answer}", self.styles['Normal']))
                        story.append(Paragraph(f"   Doğru cevap: {correct_answer}", self.styles['Normal']))
                        if explanation:
                            story.append(Paragraph(f"   Açıklama: {explanation}", self.styles['Normal']))
                        story.append(Spacer(1, 8))
                
            else:
                story.append(Paragraph("Detaylı soru bilgisi bulunamadı.", self.styles['Normal']))
            
            # Genel Değerlendirme ve Yorum
            story.append(Spacer(1, 20))
            story.append(Paragraph("Genel Değerlendirme", self.styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # Quiz sonucuna göre değerlendirme
            performance_text = ""
            if quiz_score >= 90:
                performance_text = f"""Aday quiz değerlendirmesinde çok yüksek bir başarı göstermiştir (%{quiz_score}). 
Sınava konu olan teknik yetkinliklere çok iyi düzeyde hâkimdir.
Adayın pozisyon için teknik yeterlilik açısından oldukça uygun olduğu düşünülmektedir."""
            elif quiz_score >= 75:
                performance_text = f"""Aday quiz değerlendirmesinde iyi düzeyde bir başarı göstermiştir (%{quiz_score}).
Sınava konu olan teknik yetkinliklere genel olarak hâkimdir.
Adayın pozisyon için teknik yeterlilik açısından uygun olduğu düşünülmektedir."""
            elif quiz_score >= 50:
                performance_text = f"""Aday quiz değerlendirmesinde orta düzeyde bir başarı göstermiştir (%{quiz_score}).
Sınava konu olan teknik yetkinliklerin bazılarında eksiği vardır.
Adayın ilgili konularda ek gelişim göstermesi faydalı olacaktır."""
            else:
                performance_text = f"""Aday quiz değerlendirmesinde düşük bir başarı göstermiştir (%{quiz_score}).
Sınava konu olan teknik yetkinliklerde önemli eksiklikleri vardır.
Adayın bu pozisyon için gerekli teknik yetkinlikleri geliştirmesi gerektiği düşünülmektedir."""
                
            story.append(Paragraph(performance_text, self.styles['Normal']))
            
            # Mülakat sonuçları da varsa, genel bir değerlendirme ekle
            if not self.report_data.get('quiz_only', True):
                story.append(Spacer(1, 12))
                mulakat_text = """Adayla yapılan mülakat ve teknik quiz sonuçları birlikte değerlendirildiğinde, 
adayın pozisyon için uygunluğu hakkında daha kapsamlı bir değerlendirme yapılabilir."""
                story.append(Paragraph(mulakat_text, self.styles['Normal']))
            
            # Altbilgi
            story.append(Spacer(1, 30))
            footer_style = ParagraphStyle(
                'Footer',
                parent=self.styles['Normal'],
                fontSize=10,
                textColor=colors.gray,
                alignment=1  # Ortalanmış
            )
            story.append(Paragraph("DUF Tech Mülakat Sistemi tarafından oluşturulmuştur", footer_style))
            
            # PDF oluştur
            doc.build(story)
            logger.info(f"Quiz raporu başarıyla oluşturuldu: {pdf_path}")
            return pdf_path
            
        except Exception as e:
            logger.error(f"Quiz raporu PDF oluşturma hatası: {str(e)}")
            return None

@app.route('/create_quiz', methods=['POST'])
def create_quiz():
    """Çoktan seçmeli sınav oluştur"""
    try:
        data = request.get_json()
        position = data.get('position', '')
        cv_context = data.get('cv_context', '')
        interview_code = data.get('code')
        quiz_type = data.get('quiz_type', 'technical') # technical, personality, both
        
        if not interview_code:
            interview_code = generate_interview_code()
            
        logger.info(f"Quiz oluşturuluyor: {interview_code} / {position}")
        
        # OpenAI'yi yapılandır
        client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        
        # Teknik ya da kişilik sorularına göre prompt'u ayarla
        if quiz_type == 'technical':
            quiz_prompt = f"""
                Verilen pozisyon için 20 adet çoktan seçmeli teknik soru oluştur. 
                Sorular adayın teknik bilgisini ölçmeye yönelik olmalı.
                
                Pozisyon: {position}
                
                {cv_context if cv_context else ''}
                
                Her soru şu formatta olmalı:
                {{
                  "id": "benzersiz numara",
                  "question": "soru metni",
                  "options": {{
                    "A": "seçenek A",
                    "B": "seçenek B",
                    "C": "seçenek C",
                    "D": "seçenek D"
                  }},
                  "correct_answer": "doğru cevabın harfi (A, B, C, D formatında)",
                  "explanation": "doğru cevabın açıklaması"
                }}
                
                Yanıt yalnızca JSON formatında soruların dizisi olmalıdır. Başka bir şey yazma.
            """
        elif quiz_type == 'personality':
            quiz_prompt = f"""
                Verilen pozisyon için 20 adet çoktan seçmeli kişilik/davranış sorusu oluştur.
                Sorular adayın iş ortamındaki davranışlarını, problem çözme becerilerini ve kişilik özelliklerini ölçmeye yönelik olmalı.
                
                Pozisyon: {position}
                
                {cv_context if cv_context else ''}
                
                Her soru şu formatta olmalı:
                {{
                  "id": "benzersiz numara",
                  "question": "soru metni",
                  "options": {{
                    "A": "seçenek A",
                    "B": "seçenek B",
                    "C": "seçenek C", 
                    "D": "seçenek D"
                  }},
                  "correct_answer": "doğru cevabın harfi (A, B, C veya D)",
                  "explanation": "doğru cevabın açıklaması"
                }}
                
                Yanıt yalnızca JSON formatında soruların dizisi olmalıdır. Başka bir şey yazma.
            """
        else: # both
            quiz_prompt = f"""
                Verilen pozisyon için 10 adet çoktan seçmeli teknik soru ve 10 adet çoktan seçmeli kişilik/davranış sorusu oluştur.
                Teknik sorular adayın yetkinliklerini ölçmeye yönelik olmalı.
                Kişilik soruları adayın iş ortamındaki davranışlarını ve kişilik özelliklerini ölçmeye yönelik olmalı.
                
                Pozisyon: {position}
                
                {cv_context if cv_context else ''}
                
                Her soru şu formatta olmalı:
                {{
                  "id": "benzersiz numara",
                  "question": "soru metni",
                  "options": {{
                    "A": "seçenek A",
                    "B": "seçenek B", 
                    "C": "seçenek C",
                    "D": "seçenek D"
                  }},
                  "correct_answer": "doğru cevabın harfi (sadece A, B, C veya D olmalıdır)",
                  "explanation": "doğru cevabın açıklaması"
                }}
                
                Yanıt yalnızca JSON formatında soruların dizisi olmalıdır. Başka bir şey yazma.
            """
        
        # GPT-4 ile sorular oluştur
        completion = client.chat.completions.create(
            model="gpt-4-turbo",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "Sen bir çoktan seçmeli sınav oluşturma uzmanısın. Belirtilen formatta sorular oluştur."},
                {"role": "user", "content": quiz_prompt}
            ]
        )
        
        # GPT'nin cevabını parse et
        quiz_response = json.loads(completion.choices[0].message.content)
        questions = quiz_response.get('questions', []) if 'questions' in quiz_response else []
        
        # Eğer sorular doğrudan bir liste olarak dönmüşse
        if not questions and isinstance(quiz_response, list):
            questions = quiz_response
            
        # Eğer hiç soru döndürülmediyse hata ver
        if not questions:
            logger.error("Quiz oluşturulamadı: GPT-4 cevabından sorular çıkarılamadı")
            logger.error(f"GPT-4 cevabı: {completion.choices[0].message.content}")
            return jsonify({
                'success': False,
                'error': "Quiz oluşturulurken bir hata oluştu. Lütfen tekrar deneyin."
            })
        
        # Her sorunun doğru formatta olmasını sağla
        for i, question in enumerate(questions):
            # ID'nin string olmasını sağla
            question['id'] = str(question.get('id', i + 1))
            
            # Doğru cevap formatını kontrol et
            correct_answer = question.get('correct_answer')
            if correct_answer:
                # Doğru cevabı standartlaştır: sadece A, B, C, D kullan
                correct_answer = str(correct_answer).strip().upper()
                if correct_answer not in ['A', 'B', 'C', 'D']:
                    logger.warning(f"Geçersiz doğru cevap formatı: {correct_answer}, düzeltiliyor")
                    # Eğer seçenekler arasında varsa, ilk seçeneği doğru cevap yap
                    options = question.get('options', {})
                    if options and len(options) > 0:
                        question['correct_answer'] = list(options.keys())[0]
                    else:
                        question['correct_answer'] = 'A'
                else:
                    question['correct_answer'] = correct_answer
        
        # GPT-4'ün döndürdüğü cevaplar arasında eşit dağılım oluşturma kontrolü
        all_answers = [q.get('correct_answer', '').strip().upper() for q in questions]
        unique_answers = set(all_answers)
        
        logger.info(f"GPT-4 tarafından döndürülen cevap dağılımı: {[all_answers.count(a) for a in unique_answers]}")
        
        # Eğer tek tip bir cevap varsa (hepsi A, B, C veya D ise) veya çok dengesiz bir dağılım varsa
        if len(unique_answers) == 1 or max([all_answers.count(a) for a in unique_answers]) > len(all_answers) * 0.7:
            logger.warning(f"GPT-4 cevaplarında dengesiz dağılım tespit edildi, rastgele düzeltiliyor.")
            
            # Seçenekler listesi
            option_letters = ['A', 'B', 'C', 'D']
            
            # Her cevap seçeneği için hedef sayılar (eşit dağılım)
            target_counts = {opt: len(all_answers) // 4 for opt in option_letters}
            
            # Artık değerleri rastgele dağıt
            remaining = len(all_answers) % 4
            for i in range(remaining):
                target_counts[option_letters[i]] += 1
                
            logger.info(f"Hedef cevap dağılımı: {target_counts}")
            
            # Şu anki sayılar
            current_counts = {opt: all_answers.count(opt) for opt in option_letters}
            
            # Rastgele atama yaparken dağılımı dengele
            import random
            random.shuffle(questions)  # Soruları karıştır
            
            for question in questions:
                current_answer = question.get('correct_answer', '').strip().upper()
                
                # Eğer bu cevap türünden fazla varsa, değiştir
                if current_counts.get(current_answer, 0) > target_counts.get(current_answer, 0):
                    # Hangi seçeneklerin sayısı hedefin altında
                    need_more = [opt for opt in option_letters if current_counts.get(opt, 0) < target_counts.get(opt, 0)]
                    
                    if need_more:
                        # Rastgele bir eksik seçeneği seç
                        new_answer = random.choice(need_more)
                        
                        # Sayımları güncelle
                        current_counts[current_answer] = current_counts.get(current_answer, 0) - 1
                        current_counts[new_answer] = current_counts.get(new_answer, 0) + 1
                        
                        # Sorunun cevabını güncelle
                        question['correct_answer'] = new_answer
                        logger.info(f"Soru {question.get('id')} için cevap '{current_answer}'dan '{new_answer}'a değiştirildi")
            
            # Son dağılımı kontrol et
            updated_answers = [q.get('correct_answer', '').strip().upper() for q in questions]
            updated_distribution = {opt: updated_answers.count(opt) for opt in option_letters}
            logger.info(f"Güncellenmiş cevap dağılımı: {updated_distribution}")
        
        # Quiz dosyasını kaydet
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        quiz_filename = f"quiz_{timestamp}_{interview_code}.json"
        quiz_data = {
            "position": position,
            "interview_code": interview_code,
            "quiz_type": quiz_type,
            "cv_context": cv_context,
            "created_at": datetime.now().isoformat(),
            "questions": questions
        }
        
        # Kayıt dizini oluştur
        os.makedirs(QUESTIONS_DIR, exist_ok=True)
        
        # Quiz dosyasını kaydet
        quiz_path = os.path.join(QUESTIONS_DIR, quiz_filename)
        with open(quiz_path, 'w', encoding='utf-8') as f:
            json.dump(quiz_data, f, ensure_ascii=False, indent=2)
            
        logger.info(f"Quiz oluşturuldu: {quiz_filename}")
        
        return jsonify({
            'success': True,
            'code': interview_code,
            'message': 'Quiz oluşturuldu'
        })
        
    except Exception as e:
        logger.error(f"Quiz oluşturma hatası: {str(e)}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f"Quiz oluşturulurken bir hata oluştu: {str(e)}"
        })

@app.route('/admin')
@login_required
def admin_dashboard():
    """Admin paneli ana sayfası"""
    # Admin kontrolü
    if not session.get('is_admin'):
        flash('Bu sayfaya erişim yetkiniz yok.', 'error')
        return redirect(url_for('create_interview'))
    
    try:
        # Temel istatistikleri topla
        total_interviews = 0
        total_quizzes = 0
        total_reports = 0
        quiz_scores = []
        position_counts = {}
        quiz_distribution = [0, 0, 0, 0, 0]  # 0-20%, 21-40%, 41-60%, 61-80%, 81-100%
        
        # Mülakatları ve quiz sonuçlarını topla
        interviews = []
        quizzes = []
        reports = []
        
        # interviews klasöründeki tüm mülakat dosyalarını tara
        if os.path.exists('interviews'):
            for file in os.listdir('interviews'):
                if file.endswith('.json'):
                    interview_path = os.path.join('interviews', file)
                    with open(interview_path, 'r', encoding='utf-8') as f:
                        try:
                            interview_data = json.load(f)
                            interview_code = file.replace('.json', '')
                            
                            # Temel mülakat bilgileri
                            interview = {
                                'code': interview_code,
                                'candidate_name': interview_data.get('candidate_name', 'İsimsiz Aday'),
                                'position': interview_data.get('position', 'Belirtilmemiş'),
                                'date': datetime.fromtimestamp(os.path.getctime(interview_path)).strftime('%d.%m.%Y %H:%M'),
                                'has_quiz_results': 'quiz_results' in interview_data,
                                'completed': interview_data.get('completed', False),
                                'has_report': 'report_pdf' in interview_data
                            }
                            
                            # Pozisyon istatistikleri
                            position = interview.get('position')
                            if position:
                                position_counts[position] = position_counts.get(position, 0) + 1
                            
                            # Eğer rapor varsa
                            if interview.get('has_report'):
                                report_pdf = interview_data.get('report_pdf')
                                interview['report_url'] = f"/reports/{report_pdf}"
                                
                                # Rapor istatistikleri
                                report = {
                                    'name': report_pdf,
                                    'interview_code': interview_code,
                                    'candidate_name': interview.get('candidate_name'),
                                    'date': datetime.fromtimestamp(os.path.getctime(os.path.join('reports', report_pdf)) if os.path.exists(os.path.join('reports', report_pdf)) else time.time()).strftime('%d.%m.%Y'),
                                    'type': 'mülakat',
                                    'url': f"/reports/{report_pdf}"
                                }
                                reports.append(report)
                                total_reports += 1
                            
                            # Quiz sonuçları
                            if interview.get('has_quiz_results'):
                                quiz_result = interview_data.get('quiz_results', {})
                                score = quiz_result.get('score', {})
                                
                                quiz = {
                                    'interview_code': interview_code,
                                    'candidate_name': interview.get('candidate_name'),
                                    'position': interview.get('position'),
                                    'date': datetime.fromtimestamp(time.time() if not quiz_result.get('timestamp') else datetime.strptime(quiz_result.get('timestamp'), "%Y-%m-%d %H:%M:%S").timestamp()).strftime('%d.%m.%Y'),
                                    'correct_count': score.get('correct', 0),
                                    'total_questions': score.get('total', 0),
                                    'percentage': score.get('percentage', 0),
                                    'has_report': interview.get('has_report'),
                                    'report_url': interview.get('report_url', '')
                                }
                                
                                # Quiz istatistikleri
                                if quiz.get('total_questions', 0) > 0:
                                    quiz_scores.append(quiz.get('percentage', 0))
                                    
                                    # Quiz başarı dağılımı
                                    percentage = quiz.get('percentage', 0)
                                    if percentage <= 20:
                                        quiz_distribution[0] += 1
                                    elif percentage <= 40:
                                        quiz_distribution[1] += 1
                                    elif percentage <= 60:
                                        quiz_distribution[2] += 1
                                    elif percentage <= 80:
                                        quiz_distribution[3] += 1
                                    else:
                                        quiz_distribution[4] += 1
                                
                                quizzes.append(quiz)
                                total_quizzes += 1
                            
                            interviews.append(interview)
                            total_interviews += 1
                            
                        except Exception as e:
                            logger.error(f"Mülakat dosyası işlenirken hata: {str(e)} - Dosya: {file}")
        
        # Özel quiz raporlarını ara
        if os.path.exists('reports'):
            for file in os.listdir('reports'):
                if file.startswith('quiz_raporu_'):
                    # Bu rapor zaten eklenmiş mi kontrol et
                    report_already_added = False
                    for report in reports:
                        if report.get('name') == file:
                            report_already_added = True
                            break
                    
                    if not report_already_added:
                        report_path = os.path.join('reports', file)
                        # Rapor adından mülakat kodunu çıkar
                        parts = file.split('_')
                        if len(parts) > 2:
                            interview_code = parts[2]
                            
                            # Mülakat bilgilerini bul
                            interview_info = None
                            for interview in interviews:
                                if interview.get('code') == interview_code:
                                    interview_info = interview
                                    break
                            
                            report = {
                                'name': file,
                                'interview_code': interview_code,
                                'candidate_name': interview_info.get('candidate_name', 'İsimsiz') if interview_info else 'İsimsiz',
                                'date': datetime.fromtimestamp(os.path.getctime(report_path)).strftime('%d.%m.%Y'),
                                'type': 'quiz',
                                'url': f"/reports/{file}"
                            }
                            reports.append(report)
                            total_reports += 1
        
        # Ortalama başarı oranı
        average_success_rate = int(sum(quiz_scores) / len(quiz_scores)) if quiz_scores else 0
        
        # Sıralama: En yeni mülakatlar ve quizler üstte
        interviews.sort(key=lambda x: x.get('date', ''), reverse=True)
        quizzes.sort(key=lambda x: x.get('date', ''), reverse=True)
        reports.sort(key=lambda x: x.get('date', ''), reverse=True)
        
        # Grafikler için veri
        position_labels = list(position_counts.keys())
        position_data = list(position_counts.values())
        
        return render_template('admin.html', 
                             total_interviews=total_interviews,
                             total_quizzes=total_quizzes,
                             total_reports=total_reports,
                             average_success_rate=average_success_rate,
                             interviews=interviews,
                             quizzes=quizzes,
                             reports=reports,
                             position_labels=json.dumps(position_labels),
                             position_data=json.dumps(position_data),
                             quiz_distribution=json.dumps(quiz_distribution))
    
    except Exception as e:
        logger.error(f"Admin panel hata: {str(e)}")
        traceback.print_exc()
        return render_template('error.html', error=f"Admin paneli yüklenirken bir hata oluştu: {str(e)}")

@app.route('/admin/interview/<interview_code>')
@login_required
def admin_interview_details(interview_code):
    """Belirli bir mülakatın detaylarını görüntüle"""
    # Admin kontrolü
    if not session.get('is_admin'):
        flash('Bu sayfaya erişim yetkiniz yok.', 'error')
        return redirect(url_for('create_interview'))
    
    try:
        # Mülakat dosyasını kontrol et
        json_path = os.path.join('interviews', f'{interview_code}.json')
        if not os.path.exists(json_path):
            return render_template('error.html', error=f"Mülakat bulunamadı: {interview_code}")
        
        # Mülakat verilerini oku
        with open(json_path, 'r', encoding='utf-8') as f:
            interview_data = json.load(f)
        
        # Rapor URL'si
        report_url = None
        if 'report_pdf' in interview_data:
            report_url = f"/reports/{interview_data['report_pdf']}"
        
        # Fotoğraflar varsa
        photos = []
        photos_dir = os.path.join('interviews', f'{interview_code}_photos')
        if os.path.exists(photos_dir):
            for photo_file in os.listdir(photos_dir):
                if photo_file.endswith('.jpg') or photo_file.endswith('.jpeg') or photo_file.endswith('.png'):
                    photos.append({
                        'url': f"/{photos_dir}/{photo_file}",
                        'timestamp': datetime.fromtimestamp(os.path.getctime(os.path.join(photos_dir, photo_file))).strftime('%d.%m.%Y %H:%M:%S')
                    })
        
        # Mülakat sorularını ve cevapları formatlı hale getir
        questions = []
        for question in interview_data.get('questions', []):
            questions.append({
                'question': question.get('question', ''),
                'answer': question.get('answer', 'Cevap verilmedi'),
                'timestamp': question.get('timestamp', '')
            })
        
        # Quiz sonuçları
        quiz_results = None
        if 'quiz_results' in interview_data:
            quiz_data = interview_data['quiz_results']
            score = quiz_data.get('score', {})
            quiz_results = {
                'correct_count': score.get('correct', 0),
                'total_questions': score.get('total', 0),
                'percentage': score.get('percentage', 0),
                'timestamp': quiz_data.get('timestamp', ''),
                'result_file': quiz_data.get('result_file', '')
            }
            
            # Quiz sonuç dosyasını oku
            if quiz_results.get('result_file'):
                result_path = os.path.join(QUESTIONS_DIR, quiz_results['result_file'])
                if os.path.exists(result_path):
                    with open(result_path, 'r', encoding='utf-8') as f:
                        result_data = json.load(f)
                        quiz_results['detailed_results'] = result_data.get('detailed_results', [])
                        quiz_results['analysis'] = result_data.get('analysis', {})
        
        return render_template('admin_interview_details.html',
                             interview_code=interview_code,
                             interview=interview_data,
                             questions=questions,
                             photos=photos,
                             quiz_results=quiz_results,
                             report_url=report_url)
    
    except Exception as e:
        logger.error(f"Mülakat detayları hata: {str(e)}")
        traceback.print_exc()
        return render_template('error.html', error=f"Mülakat detayları yüklenirken bir hata oluştu: {str(e)}")

@app.route('/admin/quiz/<interview_code>')
@login_required
def admin_quiz_details(interview_code):
    """Belirli bir quizin detaylarını görüntüle"""
    # Admin kontrolü
    if not session.get('is_admin'):
        flash('Bu sayfaya erişim yetkiniz yok.', 'error')
        return redirect(url_for('create_interview'))
    
    try:
        # Mülakat dosyasını kontrol et
        json_path = os.path.join('interviews', f'{interview_code}.json')
        if not os.path.exists(json_path):
            return render_template('error.html', error=f"Mülakat bulunamadı: {interview_code}")
        
        # Mülakat verilerini oku
        with open(json_path, 'r', encoding='utf-8') as f:
            interview_data = json.load(f)
        
        # Quiz sonuçları
        if 'quiz_results' not in interview_data:
            return render_template('error.html', error=f"Bu mülakat için quiz sonucu bulunamadı: {interview_code}")
        
        quiz_data = interview_data['quiz_results']
        score = quiz_data.get('score', {})
        quiz_results = {
            'interview_code': interview_code,
            'candidate_name': interview_data.get('candidate_name', 'İsimsiz Aday'),
            'position': interview_data.get('position', 'Belirtilmemiş'),
            'correct_count': score.get('correct', 0),
            'total_questions': score.get('total', 0),
            'percentage': score.get('percentage', 0),
            'timestamp': quiz_data.get('timestamp', ''),
            'result_file': quiz_data.get('result_file', '')
        }
        
        # Quiz sonuç dosyasını oku
        detailed_results = []
        analysis = {}
        if quiz_results.get('result_file'):
            result_path = os.path.join(QUESTIONS_DIR, quiz_results['result_file'])
            if os.path.exists(result_path):
                with open(result_path, 'r', encoding='utf-8') as f:
                    result_data = json.load(f)
                    detailed_results = result_data.get('detailed_results', [])
                    analysis = result_data.get('analysis', {})
        
        # Rapor URL'si
        report_url = None
        if 'report_pdf' in interview_data:
            report_url = f"/reports/{interview_data['report_pdf']}"
        
        return render_template('admin_quiz_details.html',
                             quiz=quiz_results,
                             detailed_results=detailed_results,
                             analysis=analysis,
                             report_url=report_url)
    
    except Exception as e:
        logger.error(f"Quiz detayları hata: {str(e)}")
        traceback.print_exc()
        return render_template('error.html', error=f"Quiz detayları yüklenirken bir hata oluştu: {str(e)}")

@app.route('/reports/<path:filename>')
def get_report(filename):
    """Rapor dosyalarını servis et"""
    # Kullanıcı giriş kontrolü
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    return send_from_directory('reports', filename)

@app.route('/interviews/<path:filename>')
def get_interview_file(filename):
    """Mülakat dosyalarını (fotoğraflar vb.) servis et"""
    # Kullanıcı giriş kontrolü
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    # Dosya yolunu parçala
    parts = filename.split('/')
    if len(parts) > 1:
        directory = os.path.join('interviews', parts[0])
        file = parts[1]
        return send_from_directory(directory, file)
    else:
        return send_from_directory('interviews', filename)

@app.route('/admin/analyze_cvs', methods=['POST'])
@login_required
def analyze_multiple_cvs():
    try:
        if 'cvs' not in request.files:
            return jsonify({
                'success': False,
                'error': 'CV dosyaları yüklenmedi.'
            }), 400
        
        cvs = request.files.getlist('cvs')
        position_id = request.form.get('position_id')
        
        if not position_id:
            return jsonify({
                'success': False,
                'error': 'Pozisyon ID gerekli.'
            }), 400
        
        conn = get_db_connection()
        c = conn.cursor()
        
        results = []
        for cv in cvs:
            if cv.filename:
                try:
                    # CV'yi kaydet
                    filename = secure_filename(cv.filename)
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'cvs', filename)
                    cv.save(file_path)
                    
                    # Mülakat kodu oluştur
                    interview_code = generate_interview_code()
                    
                    # CV'yi veritabanına kaydet
                    c.execute('''
                        INSERT INTO cvs (position_id, name, email, file_path, interview_code)
                        VALUES (?, ?, ?, ?, ?)
                    ''', (position_id, filename.split('.')[0], '', file_path, interview_code))
                    
                    cv_id = c.lastrowid
                    
                    # CV analizi yap
                    analyze_cv_content(file_path, position_id, cv_id)
                    
                    results.append({
                        'filename': filename,
                        'status': 'success',
                        'interview_code': interview_code
                    })
                    
                except Exception as e:
                    results.append({
                        'filename': cv.filename,
                        'status': 'error',
                        'error': str(e)
                    })
        
        conn.commit()
        
        return jsonify({
            'success': True,
            'message': f'{len(results)} CV analiz edildi.',
            'results': results
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500
    finally:
        conn.close()

def analyze_cv_content(cv_path, position_id, cv_id):
    try:
        if not cv_path:
            raise ValueError("CV dosya yolu gerekli")
            
        if not position_id:
            raise ValueError("Pozisyon ID gerekli")
            
        if not cv_id:
            raise ValueError("CV ID gerekli")
        
        # CV'yi metne çevir
        cv_text = extract_text_from_cv(cv_path)
        
        # Pozisyon gereksinimlerini al
        conn = get_db_connection()
        c = conn.cursor()
        
        c.execute('''
            SELECT requirements, preferred_skills
            FROM positions
            WHERE id = ?
        ''', (position_id,))
        
        position = c.fetchone()
        
        if not position:
            raise Exception(f"Pozisyon bulunamadı (ID: {position_id})")
        
        # CV analizi yap
        analysis = analyze_cv(cv_text, position['requirements'], position['preferred_skills'])
        
        # Sonuçları kaydet
        c.execute('''
            UPDATE cvs
            SET match_score = ?, analysis_data = ?
            WHERE id = ?
        ''', (analysis['match_score'], json.dumps(analysis), cv_id))
        
        conn.commit()
        
    except Exception as e:
        print(f"CV analizi sırasında hata: {str(e)} - Dosya: {cv_path}")
        # Hata durumunda varsayılan değerleri kaydet
        try:
            c.execute('''
                UPDATE cvs
                SET match_score = 0, analysis_data = ?
                WHERE id = ?
            ''', (json.dumps({
                'match_score': 0,
                'analysis': f"Analiz sırasında bir hata oluştu: {str(e)}"
            }), cv_id))
            conn.commit()
        except Exception as update_error:
            print(f"Varsayılan değerler kaydedilirken hata: {str(update_error)}")
    finally:
        conn.close()

def extract_text_from_cv(cv_path):
    # PDF'den metin çıkarma
    if cv_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(cv_path)
    # DOC/DOCX'ten metin çıkarma
    elif cv_path.lower().endswith(('.doc', '.docx')):
        return extract_text_from_doc(cv_path)
    else:
        return ""

def extract_text_from_pdf(pdf_path):
    try:
        import PyPDF2
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return text
    except Exception as e:
        print(f"PDF metin çıkarma hatası: {str(e)}")
        return ""

def extract_text_from_doc(doc_path):
    try:
        import docx
        doc = docx.Document(doc_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"DOC metin çıkarma hatası: {str(e)}")
        return ""

def analyze_cv(cv_text, requirements, preferred_skills):
    try:
        # OpenAI API'yi kullanarak CV analizi yap
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Sen bir CV analiz uzmanısın. Verilen CV'yi pozisyon gereksinimlerine göre analiz et."},
                {"role": "user", "content": f"""
                CV Metni:
                {cv_text}
                
                Pozisyon Gereksinimleri:
                {requirements}
                
                Tercih Edilen Özellikler:
                {preferred_skills}
                
                Lütfen bu CV'yi analiz et ve aşağıdaki formatta bir yanıt ver:
                1. Uyum Skoru (0-100 arası)
                2. Güçlü Yönler
                3. Gelişim Alanları
                4. Detaylı Değerlendirme
                """}
            ]
        )
        
        analysis_text = response.choices[0].message.content
        
        # Uyum skorunu çıkar
        match_score = extract_score(analysis_text, "Uyum Skoru")
        
        return {
            'match_score': match_score,
            'analysis': analysis_text
        }
        
    except Exception as e:
        print(f"CV analiz hatası: {str(e)}")
        return {
            'match_score': 0,
            'analysis': "Analiz sırasında bir hata oluştu."
        }

@app.route('/positions')
@login_required
def positions():
    return render_template('positions.html')

if __name__ == '__main__':
    try:
        print("\n=== DUF Tech Mülakat Sistemi Başlatılıyor ===")
        
        # Dosya izleme sistemini başlat
        observer = start_file_watcher()
        
        # Flask uygulamasını başlat
        print("\n[*] Web sunucusu başlatılıyor (Port: 5004)...")
        app.run(host='0.0.0.0', port=5006)
        
    except Exception as e:
        print(f"\n[!] Program başlatılamadı: {str(e)}")
    finally:
        if 'observer' in locals() and observer:
            observer.stop()
            observer.join()
            